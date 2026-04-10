# blueprints/planos.py

import os
import requests 
import json     
from datetime import date, datetime 
from io import BytesIO

from flask import (
    Blueprint, render_template, redirect, url_for, 
    send_file, flash, jsonify, send_from_directory, request, current_app
)
from werkzeug.utils import secure_filename
from sqlalchemy import func, case 

# --- Imports para Exportação e Documentos ---
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER
from reportlab.lib.styles import getSampleStyleSheet # Necessário para o PDF

# --- Imports de Módulos Locais ---
from app.models import (
    db, Turma, Aluno, Atividade, Presenca, PlanoDeAula, 
    Material, Horario, BlocoAula, DiarioBordo, Lembrete
)
from app.forms.forms_legacy import (
    PlanoDeAulaForm, MaterialForm, DiarioForm 
)
# Assumindo que essas funções estão em 'utils.py'
from app.utils.helpers import extrair_texto_de_ficheiro, obter_resumo_ia 
from flask_login import login_required, current_user

# Criação do Blueprint para Planejamento, Diário e Horário
planos_bp = Blueprint('planos', __name__, url_prefix='/')


# ------------------- FUNÇÕES HELPER LOCAIS -------------------

def format_text_for_pdf(text):
    """Formata texto multilinha para exibição em PDF (ReportLab)."""
    if text:
        return text.replace('\n', '<br/>')
    return "N/A"

def calcular_media_desempenho_turma(id_turma):
    """Calcula a média de desempenho da turma com base no percentual (0-100%)."""
    alunos = Aluno.query.filter_by(id_turma=id_turma).all()
    ids_alunos_turma = [aluno.id for aluno in alunos]
    
    if not ids_alunos_turma:
        return 0.0

    # O desempenho é calculado com base no campo 'desempenho' (0-100%)
    desempenho_turmas_raw = db.session.query(
        func.avg(Presenca.desempenho).label('media_desempenho')
    ).filter(Presenca.id_aluno.in_(ids_alunos_turma)).scalar()
    
    return float(desempenho_turmas_raw) if desempenho_turmas_raw else 0.0


# ------------------- ROTAS DE PLANEJAMENTO -------------------

@planos_bp.route('/planejamentos')
@login_required
def planejamentos():
    turmas = Turma.query.filter_by(autor=current_user).order_by(Turma.nome).all()
    # CORREÇÃO: Template na pasta 'geral'
    return render_template('professor/planejamento/lista.html', turmas=turmas)


@planos_bp.route('/turma/<int:id_turma>/planejamento', methods=['GET', 'POST'])
@login_required
def planejamento(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    form = PlanoDeAulaForm() 
    material_form = MaterialForm() 
    
    if form.validate_on_submit() and request.form.get('submit_manual'):
        novo_plano = PlanoDeAula(
            id_turma=id_turma,
            data_prevista=form.data_prevista.data,
            titulo=form.titulo.data,
            conteudo=form.conteudo.data,
            habilidades_bncc=form.habilidades_bncc.data,
            objetivos=form.objetivos.data,
            duracao=form.duracao.data,
            recursos=form.recursos.data,
            metodologia=form.metodologia.data,
            avaliacao=form.avaliacao.data,
            referencias=form.referencias.data,
            status='Planejado'
        )
        db.session.add(novo_plano)
        db.session.commit()
        flash('Novo plano de aula salvo com sucesso!', 'success')
        return redirect(url_for('planos.planejamento', id_turma=id_turma))
            
    planos = PlanoDeAula.query.filter_by(id_turma=id_turma).order_by(PlanoDeAula.data_prevista.desc()).all()
    
    # CORREÇÃO: Template na pasta 'geral'
    return render_template('professor/planejamento/editor.html', 
                           turma=turma, 
                           form=form, 
                           planos=planos,
                           material_form=material_form 
                           )


# --- ROTA DE IA (Gerar Plano) ---
@planos_bp.route('/turma/<int:id_turma>/gerar_plano_ia', methods=['POST'])
@login_required
def gerar_plano_ia(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('planos.planejamento', id_turma=id_turma))
    
    planos = PlanoDeAula.query.filter_by(id_turma=id_turma).order_by(PlanoDeAula.data_prevista.desc()).all()
    tema_ia = request.form.get('tema_ia')
    
    api_key = current_app.config.get('GOOGLE_API_KEY')
    if not api_key:
        flash('A chave da API do Google AI não está configurada.', 'danger')
        return redirect(url_for('planos.planejamento', id_turma=id_turma))

    prompt = f"""
    Aja como um professor especialista em pedagogia.
    Gere um plano de aula detalhado para a turma '{turma.nome}'.
    O tema da aula é: '{tema_ia}'.
    
    Por favor, responda **APENAS** com um objeto JSON, sem nenhum texto antes ou depois (sem markdown). 
    O JSON deve ter as seguintes chaves (use strings vazias "" se não houver conteúdo para uma chave):
    - "titulo": (string, o tema da aula)
    - "conteudo": (string, tópicos principais)
    - "habilidades_bncc": (string, códigos de habilidades da BNCC)
    - "objetivos": (string, o que os alunos devem ser capazes de fazer)
    - "duracao": (string, ex: "2 aulas de 50 min")
    - "recursos": (string, lista de materiais)
    - "metodologia": (string, passo-a-passo da aula)
    - "avaliacao": (string, como o aprendizado será verificado)
    - "referencias": (string, sugestões de leitura)
    """

    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    
    data = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {
            "responseMimeType": "application/json",
            "temperature": 0.7,
            "maxOutputTokens": 8192,
        }
    }
    
    dados_ia_bruto = "" 
    try:
        response = requests.post(url, headers=headers, data=json.dumps(data), timeout=20)
        response.raise_for_status() 
        
        dados_ia_bruto = response.json()['candidates'][0]['content']['parts'][0]['text']
        
        if isinstance(dados_ia_bruto, str):
            texto_json_limpo = dados_ia_bruto.strip()
            # Remove blocos de código markdown se existirem
            if texto_json_limpo.startswith('```json'):
                texto_json_limpo = texto_json_limpo[7:]
            if texto_json_limpo.endswith('```'):
                texto_json_limpo = texto_json_limpo[:-3]
                
            dados_ia = json.loads(texto_json_limpo)
        else:
            dados_ia = dados_ia_bruto 

        form = PlanoDeAulaForm()
        
        form.data_prevista.data = date.today()
        form.titulo.data = dados_ia.get('titulo', tema_ia)
        form.conteudo.data = dados_ia.get('conteudo', '')
        form.habilidades_bncc.data = dados_ia.get('habilidades_bncc', '')
        form.objetivos.data = dados_ia.get('objetivos', '')
        form.duracao.data = dados_ia.get('duracao', '')
        form.recursos.data = dados_ia.get('recursos', '')
        form.metodologia.data = dados_ia.get('metodologia', '')
        form.avaliacao.data = dados_ia.get('avaliacao', '')
        form.referencias.data = dados_ia.get('referencias', '')
        
        flash('Plano de aula gerado pela IA! Revise os dados e clique em Salvar.', 'success')

    except requests.exceptions.RequestException as e:
        flash(f'Erro ao conectar com a API de IA: {e}', 'danger')
        form = PlanoDeAulaForm() 
    except (KeyError, IndexError, json.JSONDecodeError) as e:
        flash(f'Erro ao analisar a resposta da IA. Tente novamente. Resposta: {dados_ia_bruto}', 'danger')
        form = PlanoDeAulaForm() 

    material_form = MaterialForm()
    # CORREÇÃO: Template na pasta 'geral'
    return render_template('professor/planejamento/editor.html', 
                           turma=turma, 
                           form=form, 
                           planos=planos,
                           material_form=material_form 
                           )

@planos_bp.route('/plano/<int:id_plano>/lancar', methods=['POST'])
@login_required
def lancar_atividade(id_plano):
    plano = PlanoDeAula.query.get_or_404(id_plano)
    if plano.turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    if plano.id_atividade_gerada:
        flash('Este plano já foi lançado como uma atividade.', 'danger')
        return redirect(url_for('planos.planejamento', id_turma=plano.id_turma))
        
    descricao_completa = f"""
**Plano de Aula:** {plano.titulo}
**Objetivos:** {plano.objetivos or 'N/A'}
**Conteúdo:** {plano.conteudo or 'N/A'}
**Habilidades BNCC:** {plano.habilidades_bncc or 'N/A'}
    """
    
    nova_atividade = Atividade(
        id_turma=plano.id_turma,
        titulo=plano.titulo,
        data=plano.data_prevista,
        peso=10.0, # Valor padrão para novas atividades, pode ser ajustado
        descricao=descricao_completa.strip()
    )
    db.session.add(nova_atividade)
    db.session.flush() 
    
    plano.status = 'Concluido'
    plano.id_atividade_gerada = nova_atividade.id
    
    texto_lembrete = f"Preparar próxima aula para '{plano.turma.nome}'. (Última aula: {plano.titulo})"
    novo_lembrete = Lembrete(
        texto=texto_lembrete,
        autor=current_user
    )
    db.session.add(novo_lembrete)
    
    nova_entrada_diario = DiarioBordo(
        id_user=current_user.id,
        id_turma=plano.id_turma,
        data=date.today(),
        anotacao=f"Aula dada (plano): {plano.titulo}.\n\n(Escreva sua reflexão aqui e anexe a atividade aplicada...)"
    )
    db.session.add(nova_entrada_diario)
    
    db.session.commit()
    
    flash(f'Atividade "{plano.titulo}" lançada, lembrete criado e entrada no diário registrada!', 'success')
    return redirect(url_for('planos.planejamento', id_turma=plano.id_turma))


# ------------------- ROTAS DE MATERIAIS DE AULA (CRUD) -------------------

@planos_bp.route('/plano/<int:id_plano>/add_material', methods=['POST'])
@login_required
def add_material(id_plano):
    plano = PlanoDeAula.query.get_or_404(id_plano)
    if plano.turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('planos.planejamento', id_turma=plano.id_turma))

    arquivo = request.files.get('arquivo_upload')
    link = request.form.get('link_externo')
    nome_link = request.form.get('nome_link')

    try:
        if arquivo and arquivo.filename != '':
            ext = arquivo.filename.split('.')[-1].lower()
            if ext not in ['pdf', 'docx', 'pptx', 'jpg', 'png', 'txt', 'zip', 'xls', 'xlsx']:
                flash('Tipo de ficheiro não suportado!', 'danger')
                return redirect(url_for('planos.planejamento', id_turma=plano.id_turma))

            # --- CORREÇÃO: Salvar na pasta 'docs' ---
            docs_folder = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs')
            if not os.path.exists(docs_folder):
                os.makedirs(docs_folder)

            filename_seguro = secure_filename(arquivo.filename)
            base, ext_seguro = os.path.splitext(filename_seguro)
            filename_final = f"{base}_{int(datetime.now().timestamp())}{ext_seguro}"
            
            filepath = os.path.join(docs_folder, filename_final)
            arquivo.save(filepath)
            
            novo_material = Material(id_plano_aula=id_plano, nome_arquivo=arquivo.filename, path_arquivo=filename_final)
            db.session.add(novo_material)
            flash('Ficheiro enviado com sucesso!', 'success')
            
        elif link and nome_link:
            if not (link.startswith('http://') or link.startswith('https://')):
                link = 'http://' + link
            novo_material = Material(id_plano_aula=id_plano, link_externo=link, nome_link=nome_link)
            db.session.add(novo_material)
            flash('Link adicionado com sucesso!', 'success')
        
        else:
            flash('Precisa de enviar um ficheiro ou preencher um Link e Nome do Link.', 'danger')

        db.session.commit()
    except Exception as e:
        db.session.rollback()
        flash(f'Erro ao salvar material: {e}', 'danger')

    return redirect(url_for('planos.planejamento', id_turma=plano.id_turma))


@planos_bp.route('/material/<int:id_material>/deletar', methods=['POST'])
@login_required
def deletar_material(id_material):
    material = Material.query.get_or_404(id_material)
    if material.plano_de_aula.turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    id_turma = material.plano_de_aula.id_turma
    
    try:
        if material.path_arquivo:
            # --- CORREÇÃO: Procurar na pasta 'docs' ---
            docs_folder = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs')
            filepath = os.path.join(docs_folder, material.path_arquivo)
            
            if os.path.exists(filepath):
                os.remove(filepath)
            else:
                # Fallback para arquivos antigos na raiz
                filepath_root = os.path.join(current_app.config['UPLOAD_FOLDER'], material.path_arquivo)
                if os.path.exists(filepath_root):
                    os.remove(filepath_root)
        
        db.session.delete(material)
        db.session.commit()
        flash('Material deletado com sucesso.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erro ao deletar material: {e}', 'danger')

    return redirect(url_for('planos.planejamento', id_turma=id_turma))

# --- ROTA DELETAR ATIVIDADE ---

@planos_bp.route('/delete_atividade/<int:atividade_id>', methods=['POST'])
@login_required
def delete_atividade(atividade_id):
    atividade = Atividade.query.get_or_404(atividade_id)
    
    if atividade.turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    id_turma = atividade.id_turma
    
    try:
        # 1. Apagar ficheiro associado, se houver
        if atividade.path_arquivo_anexo:
            # --- CORREÇÃO: Procurar na pasta 'docs' ---
            filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs', atividade.path_arquivo_anexo)
            if os.path.exists(filepath):
                os.remove(filepath)
            else:
                # Fallback
                filepath_root = os.path.join(current_app.config['UPLOAD_FOLDER'], atividade.path_arquivo_anexo)
                if os.path.exists(filepath_root):
                    os.remove(filepath_root)
                
        # 2. Deletar a Atividade (e suas Presenças via cascade)
        db.session.delete(atividade)
        db.session.commit()
        
        flash(f'Atividade "{atividade.titulo}" deletada com sucesso.', 'success')
        
    except Exception as e:
        db.session.rollback()
        flash(f'Erro ao deletar atividade: {e}', 'danger')
        
    return redirect(url_for('alunos.turma', id_turma=id_turma)) 


# ------------------- ROTAS DE DIÁRIO DE BORDO -------------------

@planos_bp.route('/diario', methods=['GET', 'POST'])
@login_required
def diario_bordo():
    id_turma_filtro = request.args.get('id_turma', type=int)
    
    form = DiarioForm()
    turmas_user = Turma.query.filter_by(autor=current_user).order_by(Turma.nome).all()

    if form.validate_on_submit():
        id_turma_form = request.form.get('id_turma')
        id_turma_valido = None
        
        if id_turma_form and id_turma_form != 'None':
            turma_selecionada = Turma.query.get(int(id_turma_form))
            if turma_selecionada and turma_selecionada.autor == current_user:
                id_turma_valido = turma_selecionada.id
        
        nova_entrada = DiarioBordo(
            id_user=current_user.id,
            id_turma=id_turma_valido,
            data=form.data.data,
            anotacao=form.anotacao.data
        )
        
        arquivo = request.files.get('arquivo_anexo')
        if arquivo and arquivo.filename != '':
            # --- CORREÇÃO: Salvar na pasta 'docs' ---
            docs_folder = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs')
            if not os.path.exists(docs_folder):
                os.makedirs(docs_folder)

            filename_seguro = secure_filename(arquivo.filename)
            base, ext_seguro = os.path.splitext(filename_seguro)
            filename_final = f"diario_{current_user.id}_{int(datetime.now().timestamp())}{ext_seguro}"
            
            filepath = os.path.join(docs_folder, filename_final)
            arquivo.save(filepath)
            
            nova_entrada.nome_arquivo_anexo = arquivo.filename
            nova_entrada.path_arquivo_anexo = filename_final
        
        db.session.add(nova_entrada)
        db.session.commit()
        return redirect(url_for('planos.diario_bordo', id_turma=(id_turma_valido if id_turma_valido else '')))

    query = DiarioBordo.query.filter_by(autor_diario=current_user).order_by(DiarioBordo.data.desc())
    if id_turma_filtro:
        query = query.filter_by(id_turma=id_turma_filtro)
    
    entradas = query.all()
    
    # CORREÇÃO: Template na pasta 'geral'
    return render_template('professor/turma/diario_classe.html', 
                           form=form, 
                           turmas_user=turmas_user,
                           entradas=entradas,
                           id_turma_filtro=id_turma_filtro)


@planos_bp.route('/diario_anexo/<path:filename>')
@login_required
def download_diario_anexo(filename):
    entrada = DiarioBordo.query.filter_by(path_arquivo_anexo=filename).first_or_404()
    if entrada.autor_diario != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('planos.diario_bordo'))
    
    # --- CORREÇÃO: Servir da pasta 'docs' (ou raiz como fallback) ---
    docs_folder = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs')
    directory = docs_folder if os.path.exists(os.path.join(docs_folder, filename)) else current_app.config['UPLOAD_FOLDER']

    return send_from_directory(
        directory, 
        entrada.path_arquivo_anexo, 
        as_attachment=True,
        download_name=entrada.nome_arquivo_anexo
    )


# ------------------- ROTAS DE HORÁRIO -------------------

@planos_bp.route('/gerenciar_horario')
@login_required
def gerenciar_horario():
    horario = Horario.query.filter_by(autor=current_user, ativo=True).first()
    
    if not horario: 
        horario = Horario(nome=f"Horário de {current_user.username}", autor=current_user)
        db.session.add(horario)
        horarios_texto_padrao = ["13:10", "14:00", "14:50", "16:00", "16:50"]
        for dia in range(5): 
            for pos in range(1, 6):
                bloco = BlocoAula(horario=horario, dia_semana=dia, posicao_aula=pos, texto_horario=horarios_texto_padrao[pos-1])
                db.session.add(bloco)
        db.session.commit()
    
    blocos_db = BlocoAula.query.filter_by(id_horario=horario.id).all()
    blocos_map = { (b.dia_semana, b.posicao_aula): b for b in blocos_db }
    
    turmas_user = Turma.query.filter_by(autor=current_user).order_by(Turma.nome).all()
    
    horarios_texto_raw = sorted(list(set(
        b.texto_horario for b in blocos_db if b.texto_horario
    )), key=lambda x: int(x.split(':')[0] + x.split(':')[1]))
    
    horarios_texto = horarios_texto_raw[:5] 
    if len(horarios_texto) < 5:
        horarios_texto.extend(["--:--"] * (5 - len(horarios_texto)))

    dias_semana = ["Segunda", "Terça", "Quarta", "Quinta", "Sexta"]
    
    # CORREÇÃO: Template na pasta 'geral'
    return render_template('admin/configuracoes/horario.html', 
                           horario=horario,
                           blocos_map=blocos_map,
                           turmas_user=turmas_user,
                           horarios_texto=horarios_texto,
                           dias_semana=dias_semana)

@planos_bp.route('/horario/salvar_bloco', methods=['POST'])
@login_required
def salvar_bloco_horario():
    data = request.json
    id_bloco = data.get('id_bloco')
    id_turma = data.get('id_turma')
    texto_alternativo = data.get('texto_alternativo')

    bloco = BlocoAula.query.get_or_404(id_bloco)
    if bloco.horario.autor != current_user:
        return jsonify({"status": "error", "message": "Não autorizado"}), 403

    try:
        nome_display = "VAGO"
        if id_turma and id_turma != 'None' and id_turma != '':
            turma_selecionada = Turma.query.get(int(id_turma))
            if turma_selecionada and turma_selecionada.autor == current_user:
                bloco.id_turma = int(id_turma)
                bloco.texto_alternativo = None
                nome_display = turma_selecionada.nome
            else:
                raise Exception("Turma não autorizada")
        else:
            bloco.id_turma = None
            bloco.texto_alternativo = texto_alternativo if texto_alternativo else None
            if bloco.texto_alternativo:
                nome_display = bloco.texto_alternativo

        db.session.commit()
        return jsonify({"status": "success", "nome_display": nome_display}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500


# ------------------- ROTAS DE IA E GERAÇÃO DE DOCUMENTOS -------------------

@planos_bp.route('/plano/<int:id_plano>/analisar_acessibilidade', methods=['POST'])
@login_required
def analisar_acessibilidade_ia(id_plano):
    plano = PlanoDeAula.query.get_or_404(id_plano)
    if plano.turma.autor != current_user:
        return jsonify({"status": "error", "message": "Não autorizado"}), 403

    api_key = current_app.config.get('GOOGLE_API_KEY')
    if not api_key:
        return jsonify({"status": "error", "message": "API Key não configurada."}), 500

    data = request.json
    necessidade = data.get('necessidade', 'dificuldades de aprendizagem gerais')

    prompt = f"""
    Aja como um especialista em design instrucional e educação inclusiva.
    Analise o seguinte plano de aula:
    - Tema: {plano.titulo}
    - Objetivos: {plano.objetivos}
    - Metodologia: {plano.metodologia}
    - Recursos: {plano.recursos}
    - Avaliação: {plano.avaliacao}

    Forneça 3 sugestões práticas e concisas de adaptação para '{necessidade}'.
    Responda APENAS com as sugestões (use \n para listas).
    """

    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    payload = {"contents": [{"parts": [{"text": prompt}]}]}

    try:
        response = requests.post(url, headers=headers, data=json.dumps(payload), timeout=20)
        response.raise_for_status() 
        texto_analise = response.json()['candidates'][0]['content']['parts'][0]['text']
        return jsonify({"status": "success", "analise": texto_analise.strip()})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@planos_bp.route('/gerar_prova', methods=['GET'])
@login_required
def gerar_prova():
    turmas = Turma.query.filter_by(autor=current_user).order_by(Turma.nome).all()
    # CORREÇÃO: Template na pasta 'geral'
    return render_template('professor/atividades/gerador_provas.html', turmas=turmas)

@planos_bp.route('/api/fontes_turma/<int:id_turma>')
@login_required
def api_fontes_turma(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        return jsonify({"error": "Não autorizado"}), 403
        
    planos = PlanoDeAula.query.filter_by(id_turma=id_turma).order_by(PlanoDeAula.data_prevista.desc()).all()
    atividades = Atividade.query.filter_by(id_turma=id_turma).order_by(Atividade.data.desc()).all()
    
    planos_json = [{"id": p.id, "titulo": p.titulo, "data": p.data_prevista.strftime('%d/%m/%Y')} for p in planos]
    atividades_json = [{"id": a.id, "titulo": a.titulo, "data": a.data.strftime('%d/%m/%Y') if a.data else ''} for a in atividades]
    
    return jsonify({
        "planos": planos_json,
        "atividades": atividades_json
    })


@planos_bp.route('/gerar_prova_docx', methods=['POST'])
@login_required
def gerar_prova_docx():
    # 1. Obter dados do formulário
    id_turma = request.form.get('turma')
    planos_ids = request.form.getlist('planos_ids')
    atividades_ids = request.form.getlist('atividades_ids')
    fontes_externas = request.files.getlist('fontes_externas')
    instrucoes_prova = request.form.get('instrucoes_prova')

    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('planos.gerar_prova'))

    api_key = current_app.config.get('GOOGLE_API_KEY')
    if not api_key:
        flash('A chave da API do Google AI não está configurada.', 'danger')
        return redirect(url_for('planos.gerar_prova'))

    # Cálculo da Média da Turma
    atividades_turma_max_score = Turma.query.get_or_404(id_turma).atividades
    total_max_score = sum(a.peso for a in atividades_turma_max_score if a.peso is not None)
    
    desempenho_medio_turma = calcular_media_desempenho_turma(id_turma)

    # 2. Construir o "Mega-Prompt" (AGORA COM RESUMOS)
    prompt_contexto = "--- INÍCIO DO CONTEXTO DA AULA (Resumido pela IA) ---\n"
    
    # 2a. Dados do DB (Planos)
    if planos_ids:
        planos = PlanoDeAula.query.filter(PlanoDeAula.id.in_(planos_ids)).all()
        for plano in planos:
            if plano.turma.autor == current_user: # Segurança
                texto_fonte = f"Plano: {plano.titulo}\nConteúdo: {plano.conteudo}\nObjetivos: {plano.objetivos}"
                resumo = obter_resumo_ia(texto_fonte, api_key, f"Plano de Aula '{plano.titulo}'")
                prompt_contexto += f"\nRESUMO DO PLANO '{plano.titulo}':\n{resumo}\n"

    # 2b. Dados do DB (Atividades)
    if atividades_ids:
        atividades = Atividade.query.filter(Atividade.id.in_(atividades_ids)).all()
        for atividade in atividades:
            if atividade.turma.autor == current_user: # Segurança
                texto_fonte = f"Atividade: {atividade.titulo}\nDescrição: {atividade.descricao or 'N/A'}"
                
                if atividade.path_arquivo_anexo:
                    # --- CORREÇÃO: Caminho correto para docs ---
                    # Tenta primeiro em docs, depois na raiz
                    docs_folder = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs')
                    filepath = os.path.join(docs_folder, atividade.path_arquivo_anexo)
                    
                    if not os.path.exists(filepath):
                        filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], atividade.path_arquivo_anexo)

                    if os.path.exists(filepath):
                        texto_extraido = extrair_texto_de_ficheiro(filepath, atividade.nome_arquivo_anexo)
                        texto_fonte += f"\nConteúdo do Anexo (Questões): {texto_extraido[:2000]}..."

                resumo = obter_resumo_ia(texto_fonte, api_key, f"Atividade '{atividade.titulo}'")
                prompt_contexto += f"\nRESUMO DO PLANO '{atividade.titulo}':\n{resumo}\n"

    # 2c. Dados de Ficheiros (Upload)
    if fontes_externas:
        for ficheiro in fontes_externas:
            if ficheiro and ficheiro.filename != '':
                filename = secure_filename(ficheiro.filename)
                
                file_stream = BytesIO(ficheiro.read())
                texto_extraido = extrair_texto_de_ficheiro(file_stream, filename)
                
                if texto_extraido:
                    resumo = obter_resumo_ia(texto_extraido, api_key, f"Ficheiro Anexado '{filename}'")
                    prompt_contexto += f"\nRESUMO DO FICHEIRO '{filename}':\n{resumo}\n"

    prompt_contexto += "\n--- FIM DO CONTEXTO DA AULA ---\n"
    
    # 3. Criar o Prompt Final para a IA
    prompt_final = f"""
    Aja como um professor experiente criando uma prova de avaliação.
    
    --- CONTEXTO DA TURMA ---
    Nome da Turma: '{turma.nome}'
    Descrição da Turma (Série, Idade, Nível): '{turma.descricao or 'Nível não informado.'}'
    Desempenho Médio Atual da Turma (0-100%): {desempenho_medio_turma:.1f}%
    Pontuação Máxima Total da Turma (Total de Pontos): {total_max_score:.1f}
    --- FIM DO CONTEXTO DA TURMA ---

    Use o CONTEXTO RESUMIDO DA AULA abaixo como sua principal fonte de informação.
    
    Sua tarefa é criar uma nova prova que sintetize o material da aula, MAS que seja **adaptada para o nível e o desempenho atual da turma** descrito acima.
    (Por exemplo, se a média for baixa, foque em revisão. Se for alta, adicione desafios).
    
    Siga estas instruções específicas:
    {instrucoes_prova}
    
    Responda APENAS com o texto da prova (título, campos para nome/data, e as questões).
    Não inclua 'Aqui está a prova:' ou qualquer outro texto introdutório.
    Formate bem as questões.
    """

    # 4. Chamar a API Gemini
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    payload = {"contents": [{"parts": [{"text": prompt_final}]}]}

    try:
        response = requests.post(url, headers=headers, data=json.dumps(payload), timeout=120) 
        response.raise_for_status() 
        texto_prova = response.json()['candidates'][0]['content']['parts'][0]['text']

        # 5. Criar o DOCX
        document = Document()
        document.add_heading(f'Avaliação: {turma.nome}', level=1)
        document.add_paragraph(f"Professor(a): {current_user.username}")
        document.add_paragraph("Nome: __________________________________________________ Data: ___/___/____")
        document.add_paragraph(f"Turma: {turma.nome} ({turma.descricao or 'N/A'})")
        document.add_heading('Instruções', level=2)
        document.add_paragraph(instrucoes_prova)
        document.add_heading('Questões', level=2)
        
        for linha in texto_prova.strip().split('\n'):
            document.add_paragraph(linha)

        # 6. Enviar Ficheiro
        f = BytesIO()
        document.save(f)
        f.seek(0)
        
        return send_file(
            f, 
            download_name=f"Prova_{turma.nome.replace(' ', '_')}.docx", 
            as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        flash(f"Erro ao gerar a prova com IA: {e}", "danger")
        return redirect(url_for('planos.gerar_prova'))


@planos_bp.route('/plano/<int:id_plano>/gerar_questoes_docx')
@login_required
def gerar_questoes_docx(id_plano):
    plano = PlanoDeAula.query.get_or_404(id_plano)
    if plano.turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    id_turma = plano.id_turma 

    api_key = current_app.config.get('GOOGLE_API_KEY')
    if not api_key:
        flash('A chave da API do Google AI não está configurada.', 'danger')
        return redirect(url_for('planos.planejamento', id_turma=id_turma))

    atividades_turma_max_score = Turma.query.get_or_404(id_turma).atividades
    total_max_score = sum(a.peso for a in atividades_turma_max_score if a.peso is not None)
    
    desempenho_medio_turma = calcular_media_desempenho_turma(id_turma)

    # 1. Construir o Contexto (Plano)
    prompt_contexto = f"""
    --- INÍCIO DO CONTEXTO (Plano de Aula) ---
    Plano: {plano.titulo}
    Conteúdo: {plano.conteudo or 'N/A'}
    Objetivos: {plano.objetivos or 'N/A'}
    Metodologia: {plano.metodologia or 'N/A'}
    --- FIM DO CONTEXTO ---
    """

    # 2. Criar o Prompt Final para a IA
    prompt_final = f"""
    Aja como um professor experiente.
    
    --- CONTEXTO DA TURMA ---
    Nome da Turma: '{plano.turma.nome}'
    Descrição da Turma (Série, Idade, Nível): '{plano.turma.descricao or 'Nível não informado.'}'
    Desempenho Médio Atual da Turma (0-100%): {desempenho_medio_turma:.1f}%
    Pontuação Máxima Total da Turma (Total de Pontos): {total_max_score:.1f}
    --- FIM DO CONTEXTO DA TURMA ---

    Use o CONTEXTO DO PLANO DE AULA acima como sua única fonte de informação.
    
    Sua tarefa é criar uma lista de 10 questões de avaliação (ex: 5 abertas, 5 múltipla escolha) baseadas **estritamente** neste contexto e **adaptadas para o nível e o desempenho atual da turma** (Série, Idade, Nível) descrito acima.
    (Por exemplo, se a média for baixa, foque em revisão. Se for alta, adicione desafios).
    
    Responda APENAS com o texto das questões, formatado para um documento.
    Não inclua 'Aqui está a prova:' ou qualquer outro texto introdutório.
    """

    # 3. Chamar a API Gemini
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    payload = {"contents": [{"parts": [{"text": prompt_final}]}]}

    try:
        response = requests.post(url, headers=headers, data=json.dumps(payload), timeout=120) 
        response.raise_for_status() 
        texto_questoes = response.json()['candidates'][0]['content']['parts'][0]['text']

        # 4. Criar o DOCX
        document = Document()
        document.add_heading(f'Questões Sugeridas: {plano.titulo}', level=1)
        document.add_paragraph(f"Turma: {plano.turma.nome} ({plano.turma.descricao or 'N/A'})")
        document.add_paragraph("Professor(a): _________________")
        document.add_paragraph("Nome: __________________________________________________ Data: ___/___/____")
        document.add_heading('Questões', level=2)
        
        for linha in texto_questoes.strip().split('\n'):
            document.add_paragraph(linha)

        # 5. Enviar Ficheiro
        f = BytesIO()
        document.save(f)
        f.seek(0)
        
        return send_file(
            f, 
            download_name=f"Questoes_{plano.titulo.replace(' ', '_')}.docx", 
            as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        flash(f"Erro ao gerar as questões com IA: {e}", "danger")
        return redirect(url_for('planos.planejamento', id_turma=plano.id_turma))


@planos_bp.route('/plano/<int:id_plano>/exportar_docx')
@login_required
def exportar_docx(id_plano):
    plano = PlanoDeAula.query.get_or_404(id_plano)
    if plano.turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    try:
        document = Document()
        document.add_heading(plano.titulo, level=1)
        document.add_paragraph(f"Data Prevista: {plano.data_prevista.strftime('%d/%m/%Y') if plano.data_prevista else 'N/A'}")
        document.add_paragraph(f"Duração: {plano.duracao or 'N/A'}")

        document.add_heading('Habilidades BNCC', level=2)
        document.add_paragraph(plano.habilidades_bncc or 'N/A')
        document.add_heading('Objetivos', level=2)
        document.add_paragraph(plano.objetivos or 'N/A')
        document.add_heading('Conteúdo', level=2)
        document.add_paragraph(plano.conteudo or 'N/A')
        document.add_heading('Metodologia', level=2)
        document.add_paragraph(plano.metodologia or 'N/A')
        document.add_heading('Recursos', level=2)
        document.add_paragraph(plano.recursos or 'N/A')
        document.add_heading('Avaliação', level=2)
        document.add_paragraph(plano.avaliacao or 'N/A')
        document.add_heading('Referências', level=2)
        document.add_paragraph(plano.referencias or 'N/A')

        f = BytesIO()
        document.save(f)
        f.seek(0)
        
        return send_file(
            f, 
            download_name=f"Plano_de_Aula_{plano.titulo.replace(' ', '_')}.docx", 
            as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        flash(f"Erro ao gerar DOCX: {e}", "danger")
        return redirect(url_for('planos.planejamento', id_turma=plano.id_turma))


@planos_bp.route('/plano/<int:id_plano>/exportar_pdf')
@login_required
def exportar_pdf(id_plano):
    plano = PlanoDeAula.query.get_or_404(id_plano)
    if plano.turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    try:
        f = BytesIO()
        doc = SimpleDocTemplate(f, pagesize=A4, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
        story = []
        styles = getSampleStyleSheet()
        
        style_h1 = ParagraphStyle(name='Heading1', fontSize=16, alignment=TA_CENTER, spaceAfter=20)
        style_h2 = ParagraphStyle(name='Heading2', fontSize=12, fontName="Helvetica-Bold", spaceAfter=6, spaceBefore=12)
        style_body = ParagraphStyle(name='BodyText', fontSize=10, alignment=TA_LEFT, spaceAfter=10)
        style_body_justify = ParagraphStyle(name='BodyTextJustify', fontSize=10, alignment=TA_JUSTIFY, spaceAfter=10)

        story.append(Paragraph(plano.titulo, style_h1))
        story.append(Paragraph(f"<b>Data Prevista:</b> {plano.data_prevista.strftime('%d/%m/%Y') if plano.data_prevista else 'N/A'}", style_body))
        story.append(Paragraph(f"<b>Duração:</b> {plano.duracao or 'N/A'}", style_body))
        
        story.append(Paragraph('Habilidades BNCC', style_h2))
        story.append(Paragraph(format_text_for_pdf(plano.habilidades_bncc), style_body_justify))
        story.append(Paragraph('Objetivos', style_h2))
        story.append(Paragraph(format_text_for_pdf(plano.objetivos), style_body_justify))
        story.append(Paragraph('Conteúdo', style_h2))
        story.append(Paragraph(format_text_for_pdf(plano.conteudo), style_body_justify))
        story.append(Paragraph('Metodologia', style_h2))
        story.append(Paragraph(format_text_for_pdf(plano.metodologia), style_body_justify))
        story.append(Paragraph('Recursos', style_h2))
        story.append(Paragraph(format_text_for_pdf(plano.recursos), style_body_justify))
        story.append(Paragraph('Avaliação', style_h2))
        story.append(Paragraph(format_text_for_pdf(plano.avaliacao), style_body_justify))
        
        story.append(Paragraph('Referências', style_h2))
        story.append(Paragraph(format_text_for_pdf(plano.referencias), style_body_justify))
        
        doc.build(story)
        f.seek(0)
        
        return send_file(
            f, 
            download_name=f"Plano_de_Aula_{plano.titulo.replace(' ', '_')}.pdf", 
            as_attachment=True,
            mimetype='application/pdf'
        )
    except Exception as e:
        flash(f"Erro ao gerar PDF: {e}", "danger")
        return redirect(url_for('planos.planejamento', id_turma=plano.id_turma))


@planos_bp.route('/diario/sugerir_ia', methods=['POST'])
@login_required
def sugerir_diario_ia():
    data_json = request.json
    id_turma = data_json.get('id_turma')
    data_diario = data_json.get('data') # Formato YYYY-MM-DD

    if not id_turma or not data_diario:
        return jsonify({"status": "error", "message": "Selecione uma Turma e uma Data."}), 400

    turma = Turma.query.get(id_turma)
    if not turma or turma.autor != current_user:
        return jsonify({"status": "error", "message": "Turma inválida."}), 403

    api_key = current_app.config.get('GOOGLE_API_KEY')
    
    # 1. Buscar Contexto (O que estava planejado vs O que foi feito)
    planos = PlanoDeAula.query.filter_by(id_turma=id_turma, data_prevista=data_diario).all()
    atividades = Atividade.query.filter_by(id_turma=id_turma, data=data_diario).all()
    
    contexto = f"Turma: {turma.nome}\nData: {data_diario}\n"
    
    if planos:
        contexto += f"Planejado: {planos[0].titulo} (Conteúdo: {planos[0].conteudo})\n"
    else:
        contexto += "Sem plano de aula específico cadastrado para hoje.\n"
        
    if atividades:
        contexto += f"Atividade Aplicada: {atividades[0].titulo} (Tipo: {atividades[0].tipo if hasattr(atividades[0], 'tipo') else 'Geral'})\n"

    # 2. Construção do Prompt
    prompt = f"""
    Aja como um professor escrevendo seu Diário de Bordo (registro de classe oficial).
    Com base no contexto abaixo, escreva um relato sucinto (1 parágrafo, max 4 linhas) sobre como foi a aula.
    Mencione o conteúdo abordado e se houve atividade avaliativa. Use tom formal mas direto.
    
    CONTEXTO:
    {contexto}
    
    Texto do Diário:
    """
    
    # 3. Chamada à API
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    data = {"contents": [{"parts": [{"text": prompt}]}]}

    try:
        response = requests.post(url, headers=headers, data=json.dumps(data), timeout=30)
        response.raise_for_status()
        sugestao = response.json()['candidates'][0]['content']['parts'][0]['text']
        return jsonify({"status": "success", "sugestao": sugestao.strip()})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

# ------------------- ROTA DE LISTAGEM GERAL (NOVA) -------------------

@planos_bp.route('/listar')
@login_required
def listar_planos():
    # Recupera todos os planos de todas as turmas do usuário
    turmas_ids = [t.id for t in current_user.turmas]
    
    if not turmas_ids:
        planos = []
    else:
        # CORREÇÃO: Usar 'id_turma' (DB) e 'data_prevista'
        planos = PlanoDeAula.query.filter(PlanoDeAula.id_turma.in_(turmas_ids)).order_by(PlanoDeAula.data_prevista.desc()).all()
        
    # CORREÇÃO: Template na pasta 'list'
    return render_template('professor/planejamento/todos_planos.html', planos=planos)