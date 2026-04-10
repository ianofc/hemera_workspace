# blueprints/alunos.py

import os
import requests 
import json 
import base64    
from datetime import date, datetime 
from io import BytesIO

from flask import (
    Blueprint, render_template, redirect, url_for, 
    send_file, flash, jsonify, send_from_directory, request, current_app
)
import pandas as pd
from werkzeug.utils import secure_filename
from sqlalchemy import func, distinct, case
from sqlalchemy.orm import joinedload

# --- Imports para Exportação de Documentos ---
# PDF (ReportLab)
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet

# Word (Docx)
from docx import Document
from docx.shared import Cm, Pt, RGBColor 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# Excel (Openpyxl)
from openpyxl.styles import Font

# Imports Locais e Extensões
from app.extensions import csrf 
from app.models import (
    db, Turma, Aluno, Atividade, Presenca, DiarioBordo, Material, BlocoAula, Horario
)
from app.forms.forms_legacy import (
    AlunoForm, AtividadeForm, PresencaForm, EditarAlunoForm
)
from app.utils.helpers import extrair_texto_de_ficheiro, obter_resumo_ia, allowed_file 
from flask_login import login_required, current_user

# Criação do Blueprint
alunos_bp = Blueprint('alunos', __name__, url_prefix='/')


# ------------------- ROTAS DE TURMAS, ALUNOS E ATIVIDADES (CRUD) -------------------

@alunos_bp.route('/turma/<int:id_turma>')
@login_required 
def turma(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Acesso não autorizado a esta turma.', 'danger')
        return redirect(url_for('core.index')) 
    
    q_aluno = request.args.get('q', '') 
    
    alunos_query = Aluno.query.filter_by(id_turma=id_turma)
    if q_aluno:
        alunos_query = alunos_query.filter(Aluno.nome.ilike(f'%{q_aluno}%'))
    
    alunos = alunos_query.order_by(Aluno.nome).all()
    
    # Carrega atividades
    atividades = Atividade.query.filter_by(id_turma=id_turma).order_by(Atividade.data.desc()).all()
    
    # --- LÓGICA DE UNIDADES ---
    ordem_unidades = [
        '1ª Unidade', '2ª Unidade', '3ª Unidade', '4ª Unidade', 
        'Recuperação', 'Exame Final'
    ]
    
    atividades_por_unidade = {}
    totais_por_unidade = {}
    
    # Agrupa atividades
    for a in atividades:
        u = a.unidade if a.unidade else '1ª Unidade' # Default para 1ª se vazio
        if u not in atividades_por_unidade:
            atividades_por_unidade[u] = []
            totais_por_unidade[u] = 0.0
        atividades_por_unidade[u].append(a)
        totais_por_unidade[u] += (a.peso or 0)

    # DESCUBRA A ÚLTIMA UNIDADE ATIVA (Para mostrar no acumulado)
    unidade_focada = '1ª Unidade' # Padrão
    # Percorre a ordem de trás para frente. A primeira que tiver dados é a atual.
    for u in reversed(ordem_unidades):
        if u in atividades_por_unidade:
            unidade_focada = u
            break
            
    # Carrega presenças
    ids_alunos = [aluno.id for aluno in alunos]
    presencas_turma = []
    if ids_alunos:
        presencas_turma = Presenca.query.join(Atividade).filter(
            Atividade.id_turma == id_turma,
            Presenca.id_aluno.in_(ids_alunos)
        ).options(joinedload(Presenca.atividade)).all() 

    # Mapeia presenças
    presencas_por_aluno = {}
    for p in presencas_turma:
        if p.id_aluno not in presencas_por_aluno:
            presencas_por_aluno[p.id_aluno] = []
        presencas_por_aluno[p.id_aluno].append(p)

    total_max_score = sum(a.peso for a in atividades if a.peso is not None)

    alunos_com_media = []
    for aluno in alunos:
        presencas_aluno = presencas_por_aluno.get(aluno.id, [])
        
        # --- CÁLCULO INTELIGENTE ---
        # Soma apenas as notas da "Unidade Focada"
        total_pontos_focados = 0
        
        for p in presencas_aluno:
            # Verifica se a atividade pertence à unidade ativa
            unidade_ativ = p.atividade.unidade if p.atividade.unidade else '1ª Unidade'
            
            if unidade_ativ == unidade_focada and p.nota is not None:
                total_pontos_focados += p.nota 
        
        alunos_com_media.append({
            'aluno': aluno,
            'media': total_pontos_focados # Agora mostra só o acumulado da unidade atual
        })
        
    return render_template(
        'professor/turma/visao_geral.html', 
        turma=turma, 
        alunos_com_media=alunos_com_media, 
        atividades=atividades,
        atividades_por_unidade=atividades_por_unidade,
        totais_por_unidade=totais_por_unidade,
        ordem_unidades=ordem_unidades,
        unidade_focada=unidade_focada, # Passamos qual unidade está sendo exibida na lista
        q=q_aluno,
        total_max_score=total_max_score
    )

@alunos_bp.route('/aluno/<int:id_aluno>')
@login_required 
def aluno(id_aluno):
    aluno = Aluno.query.get_or_404(id_aluno)
    if not aluno.turma or aluno.turma.autor != current_user:
        flash('Acesso não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    # Busca atividades e presenças
    atividades = Atividade.query.filter_by(id_turma=aluno.id_turma).order_by(Atividade.data.desc()).all()
    presencas = Presenca.query.filter_by(id_aluno=id_aluno).all()
    presencas_map = {p.id_atividade: p for p in presencas}
    
    # --- ESTRUTURAÇÃO DO BOLETIM POR UNIDADE ---
    ordem_unidades = [
        '1ª Unidade', '2ª Unidade', '3ª Unidade', '4ª Unidade', 
        'Recuperação', 'Exame Final'
    ]
    
    # Dicionário mestre: { '1ª Unidade': {'atividades': [], 'total_obtido': 0, 'total_max': 0} }
    boletim = {}
    
    # Inicializa
    for u in ordem_unidades:
        boletim[u] = {'atividades': [], 'total_obtido': 0.0, 'total_max': 0.0, 'tem_dados': False}
        
    media_geral_acumulada = 0.0
    total_geral_max = 0.0

    for ativ in atividades:
        u = ativ.unidade if ativ.unidade else '1ª Unidade'
        if u not in boletim: u = '1ª Unidade' # Fallback
        
        # Dados da atividade
        p = presencas_map.get(ativ.id)
        nota = p.nota if (p and p.nota is not None) else 0.0
        status = p.status if p else 'Pendente'
        
        # Adiciona ao boletim
        boletim[u]['atividades'].append({
            'atividade': ativ,
            'presenca': p,
            'nota': nota
        })
        
        # Somatórios da Unidade
        boletim[u]['total_max'] += (ativ.peso or 0)
        boletim[u]['total_obtido'] += nota
        boletim[u]['tem_dados'] = True
        
        # Somatórios Gerais
        total_geral_max += (ativ.peso or 0)
        media_geral_acumulada += nota

    return render_template(
        'aluno/perfil.html', 
        aluno=aluno, 
        boletim=boletim, # Passamos a estrutura organizada
        ordem_unidades=ordem_unidades,
        media_geral_acumulada=media_geral_acumulada,
        total_geral_max=total_geral_max
    )

@alunos_bp.route('/add_aluno/<int:id_turma>', methods=['GET', 'POST'])
@login_required 
def add_aluno(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    form = AlunoForm()
    if form.validate_on_submit():
        novo = Aluno(
            nome=form.nome.data, 
            matricula=form.matricula.data, 
            id_turma=id_turma, 
            data_cadastro=date.today()
        )
        db.session.add(novo)
        db.session.commit()
        flash(f'Aluno {novo.nome} adicionado!', 'success')
        return redirect(url_for('alunos.turma', id_turma=id_turma))
    
    return render_template('admin/usuarios/novo_aluno.html', turma=turma, form=form)

@alunos_bp.route('/turma/<int:id_turma>/bulk_add_alunos', methods=['POST'])
@login_required
def bulk_add_alunos(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        return jsonify({"status": "error", "message": "Não autorizado"}), 403

    data = request.json
    nomes = data.get('nomes', [])

    if not nomes:
        return jsonify({"status": "error", "message": "Nenhum nome enviado"}), 400

    try:
        nomes_adicionados = 0
        for nome in nomes:
            nome_limpo = nome.strip()
            if nome_limpo: # Ignora linhas em branco
                novo_aluno = Aluno(
                    nome=nome_limpo,
                    id_turma=id_turma,
                    data_cadastro=date.today()
                )
                db.session.add(novo_aluno)
                nomes_adicionados += 1
        
        db.session.commit()
        return jsonify({"status": "success", "message": f"{nomes_adicionados} alunos adicionados."}), 200
    
    except Exception as e:
        db.session.rollback()
        print(f"Erro ao adicionar alunos em massa: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500


@alunos_bp.route('/add_atividade/<int:id_turma>', methods=['GET', 'POST'])
@login_required 
def add_atividade(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    form = AtividadeForm()
    if form.validate_on_submit():
        
        # --- Lógica de Upload do Anexo (Salva na pasta DOCS) ---
        arquivo = request.files.get('arquivo_anexo')
        nome_arquivo_anexo = None
        path_arquivo_anexo = None
        
        if arquivo and arquivo.filename != '' and allowed_file(arquivo.filename):
            # Define pasta de documentos
            docs_folder = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs')
            if not os.path.exists(docs_folder):
                os.makedirs(docs_folder) # Cria a pasta docs se não existir

            filename_seguro = secure_filename(arquivo.filename)
            _, ext_seguro = os.path.splitext(filename_seguro)
            filename_final = f"atividade_{id_turma}_{int(datetime.now().timestamp())}{ext_seguro}"
            
            # Salva em uploads/docs/
            filepath = os.path.join(docs_folder, filename_final)
            arquivo.save(filepath)
            
            nome_arquivo_anexo = arquivo.filename
            path_arquivo_anexo = filename_final
        # --- Fim da Lógica de Upload ---
        
        # Lógica para incluir número de questões na descrição
        descricao = form.descricao.data
        if form.num_questoes.data: 
            descricao = f"Nº de Questões: {form.num_questoes.data}\n\n{descricao}"

        atividade = Atividade(
            id_turma=id_turma, 
            titulo=form.titulo.data,
            unidade=form.unidade.data, # <--- SALVA A UNIDADE AQUI
            data=form.data.data, 
            peso=form.valor_total.data, 
            tipo=form.tipo.data,        
            descricao=descricao,
            nome_arquivo_anexo=nome_arquivo_anexo,
            path_arquivo_anexo=path_arquivo_anexo
        )
        
        db.session.add(atividade)
        db.session.commit()
        flash(f'Atividade criada na {atividade.unidade}!', 'success')
        return redirect(url_for('alunos.turma', id_turma=id_turma))
    
    ai_desc = request.args.get('ai_desc', None)
    if ai_desc:
        form.descricao.data = ai_desc
        
    return render_template('professor/atividades/nova_atividade.html', turma=turma, form=form)


@alunos_bp.route('/registrar_presenca/<int:id_aluno>/<int:id_atividade>', methods=['GET', 'POST'])
@login_required 
@csrf.exempt
def registrar_presenca(id_aluno, id_atividade):
    aluno = Aluno.query.get_or_404(id_aluno)
    atividade = Atividade.query.get_or_404(id_atividade)
    if aluno.turma.autor != current_user or atividade.turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    presenca = Presenca.query.filter_by(id_aluno=id_aluno, id_atividade=id_atividade).first()
    form = PresencaForm(obj=presenca) 
    
    if form.validate_on_submit():
        if form.nota.data is None:
             form.nota.data = 0.0

        # Lógica de Backup Backend: Se veio 'acertos' e 'total_questoes_manual' no request, recalcula
        total_questoes_manual = request.form.get('total_questoes_manual')
        if form.acertos.data is not None and total_questoes_manual and int(total_questoes_manual) > 0:
             acertos = form.acertos.data
             total = int(total_questoes_manual)
             if acertos <= total:
                 nota_calc = (acertos / total) * atividade.peso
                 form.nota.data = round(nota_calc, 2)

        if presenca:
            form.populate_obj(presenca)
            flash(f'Registro atualizado! Nota: {presenca.nota}', 'success')
        else:
            nova_presenca = Presenca(id_aluno=id_aluno, id_atividade=id_atividade)
            form.populate_obj(nova_presenca)
            db.session.add(nova_presenca)
            flash(f'Registro criado! Nota: {nova_presenca.nota}', 'success')
        
        db.session.commit()
        return redirect(url_for('alunos.turma', id_turma=aluno.id_turma))

    return render_template('professor/turma/chamada.html', aluno=aluno, atividade=atividade, form=form)

@alunos_bp.route('/atividade/<int:id_atividade>/editar', methods=['GET', 'POST'])
@login_required
def edit_atividade(id_atividade):
    atividade = Atividade.query.get_or_404(id_atividade)
    if atividade.turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    class TempAtividade:
        def __init__(self, atividade):
            self.__dict__ = atividade.__dict__.copy()
            self.valor_total = atividade.peso
            
            num_questoes = None
            if atividade.descricao:
                try:
                    num_questoes_match = next((line.split(':')[-1].strip() for line in atividade.descricao.split('\n') if line.strip().startswith("Nº de Questões:")), None)
                    if num_questoes_match:
                         num_questoes = int(num_questoes_match)
                except:
                    num_questoes = None
            self.num_questoes = num_questoes 

    temp_atividade = TempAtividade(atividade)
    form = AtividadeForm(obj=temp_atividade)
    
    if form.validate_on_submit():
        
        # Lógica de Upload do Anexo
        arquivo = request.files.get('arquivo_anexo')
        if arquivo and arquivo.filename != '' and allowed_file(arquivo.filename):
            docs_folder = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs')
            if not os.path.exists(docs_folder):
                os.makedirs(docs_folder)

            if atividade.path_arquivo_anexo:
                # Tenta apagar arquivo antigo
                old_path = os.path.join(docs_folder, atividade.path_arquivo_anexo)
                if os.path.exists(old_path):
                    os.remove(old_path)
                else:
                    old_path_root = os.path.join(current_app.config['UPLOAD_FOLDER'], atividade.path_arquivo_anexo)
                    if os.path.exists(old_path_root):
                        os.remove(old_path_root)
            
            filename_seguro = secure_filename(arquivo.filename)
            _, ext_seguro = os.path.splitext(filename_seguro)
            filename_final = f"atividade_{atividade.id_turma}_{int(datetime.now().timestamp())}{ext_seguro}"
            
            filepath = os.path.join(docs_folder, filename_final)
            arquivo.save(filepath)
            
            atividade.nome_arquivo_anexo = arquivo.filename
            atividade.path_arquivo_anexo = filename_final
        
        form.populate_obj(atividade)
        
        atividade.peso = form.valor_total.data
        atividade.unidade = form.unidade.data # <--- ATUALIZA A UNIDADE AQUI

        descricao_form_data = form.descricao.data

        if form.num_questoes.data:
            lines = descricao_form_data.split('\n')
            
            start_index = 0
            if lines and lines[0].strip().startswith("Nº de Questões:"):
                try:
                    for i, line in enumerate(lines):
                        if i > 0 and line.strip() != '':
                            start_index = i
                            break
                        elif i == len(lines) - 1 and line.strip() == '':
                            start_index = len(lines)
                            break
                    if start_index == 0: 
                        if len(lines) >= 3 and lines[1].strip() == '':
                            start_index = 2
                        else:
                            start_index = 1
                except Exception:
                    start_index = 1

            descricao_base_limpa = "\n".join(lines[start_index:]).strip()
            atividade.descricao = f"Nº de Questões: {form.num_questoes.data}\n\n{descricao_base_limpa}"
        else:
            lines = atividade.descricao.split('\n') if atividade.descricao else []
            if lines and lines[0].strip().startswith("Nº de Questões:"):
                atividade.descricao = "\n".join(lines[2:]).strip()
            else:
                atividade.descricao = descricao_form_data.strip() 

        if not atividade.descricao.strip():
            atividade.descricao = None
        
        db.session.commit()
        flash(f'Atividade "{atividade.titulo}" atualizada com sucesso!', 'success')
        return redirect(url_for('alunos.turma', id_turma=atividade.id_turma))

    # CORREÇÃO: Template na pasta edit/
    return render_template('edit/edit_atividade.html', form=form, atividade=atividade)


@alunos_bp.route('/atividade/<int:id_atividade>/deletar', methods=['POST'])
@login_required
def delete_atividade(id_atividade):
    atividade = Atividade.query.get_or_404(id_atividade)
    if atividade.turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))

    id_turma = atividade.id_turma
    try:
        if atividade.path_arquivo_anexo:
            filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs', atividade.path_arquivo_anexo)
            if os.path.exists(filepath):
                os.remove(filepath)
            else:
                filepath_root = os.path.join(current_app.config['UPLOAD_FOLDER'], atividade.path_arquivo_anexo)
                if os.path.exists(filepath_root):
                    os.remove(filepath_root)
        
        db.session.delete(atividade)
        db.session.commit()
        flash(f'Atividade "{atividade.titulo}" deletada.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erro ao deletar atividade: {e}', 'danger')

    return redirect(url_for('alunos.turma', id_turma=id_turma))


@alunos_bp.route('/aluno/<int:id_aluno>/editar', methods=['GET', 'POST'])
@login_required 
def editar_aluno(id_aluno):
    aluno = Aluno.query.get_or_404(id_aluno)
    if (not aluno.turma) or (aluno.turma.autor != current_user):
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    form = EditarAlunoForm(obj=aluno)
    form.turma.query_factory = lambda: Turma.query.filter_by(autor=current_user).order_by(Turma.nome)
    
    if form.validate_on_submit():
        if form.turma.data and form.turma.data.autor != current_user:
            flash('Seleção de turma inválida.', 'danger')
            return redirect(url_for('alunos.editar_aluno', id_aluno=id_aluno))
            
        form.populate_obj(aluno) 
        db.session.commit()
        return redirect(url_for('alunos.aluno', id_aluno=aluno.id))
    
    # CORREÇÃO: Template na pasta edit/
    return render_template('edit/edit_aluno.html', aluno=aluno, form=form)

@alunos_bp.route('/aluno/<int:id_aluno>/deletar', methods=['POST'])
@login_required 
def deletar_aluno(id_aluno):
    aluno = Aluno.query.get_or_404(id_aluno)
    if (not aluno.turma) or (aluno.turma.autor != current_user):
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    id_turma = aluno.id_turma
    db.session.delete(aluno)
    db.session.commit()
    flash(f'Aluno {aluno.nome} deletado com sucesso!', 'danger')
    return redirect(url_for('alunos.turma', id_turma=id_turma))

@alunos_bp.route('/turma/<int:id_turma>/editar', methods=['GET', 'POST'])
@login_required 
def editar_turma(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    # CORREÇÃO DE IMPORT LOCAL:
    from app.forms.forms_legacy import TurmaForm 
    form = TurmaForm(obj=turma)
    if form.validate_on_submit():
        form.populate_obj(turma)
        db.session.commit()
        flash(f'Turma {turma.nome} atualizada com sucesso!', 'success')
        return redirect(url_for('alunos.turma', id_turma=turma.id))
    
    # CORREÇÃO: Template na pasta edit/
    return render_template('edit/edit_turma.html', turma=turma, form=form)

@alunos_bp.route('/turma/<int:id_turma>/deletar', methods=['POST'])
@login_required 
def deletar_turma(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    Aluno.query.filter_by(id_turma=id_turma).update({'id_turma': None})
    
    user_horario_ids = [h.id for h in current_user.horarios]
    if user_horario_ids:
        BlocoAula.query.filter(
            BlocoAula.id_horario.in_(user_horario_ids),
            BlocoAula.id_turma == id_turma
        ).update({'id_turma': None, 'texto_alternativo': f'Turma Apagada ({turma.nome})'})
    
    db.session.delete(turma)
    db.session.commit()
    flash(f'Turma {turma.nome} deletada. Os alunos desta turma ficaram sem turma.', 'danger')
    return redirect(url_for('core.index'))

# ------------------- ROTAS DE GRADEBOOK E DASHBOARD POR TURMA -------------------

@alunos_bp.route('/turma/<int:id_turma>/gradebook')
@login_required
def gradebook(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Acesso não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    # 1. Definir a Lista COMPLETA de Unidades (Não depende do banco)
    lista_unidades = [
        '1ª Unidade', 
        '2ª Unidade', 
        '3ª Unidade', 
        '4ª Unidade', 
        'Recuperação', 
        'Exame Final',
        'Todas' # Opção para ver o ano todo
    ]
    
    # 2. Capturar a escolha do usuário (ou definir padrão)
    unidade_selecionada = request.args.get('unidade')
    
    # Lógica inteligente de padrão:
    # Se não escolheu nada, tenta pegar a unidade da atividade mais recente cadastrada.
    # Se não tiver atividades, vai para a '1ª Unidade'.
    if not unidade_selecionada:
        ultima_atividade = Atividade.query.filter_by(id_turma=id_turma).order_by(Atividade.data.desc()).first()
        if ultima_atividade and ultima_atividade.unidade:
            unidade_selecionada = ultima_atividade.unidade
        else:
            unidade_selecionada = '1ª Unidade'

    # 3. Buscar Alunos
    alunos = Aluno.query.filter_by(id_turma=id_turma).order_by(Aluno.nome).all()
    
    # 4. Buscar Atividades (COM FILTRO REAL)
    query_atividades = Atividade.query.filter_by(id_turma=id_turma)
    
    # Se NÃO for "Todas", aplica o filtro estrito
    if unidade_selecionada != 'Todas':
        query_atividades = query_atividades.filter_by(unidade=unidade_selecionada)
    
    # Ordena por Data (para ficar cronológico na tabela)
    atividades = query_atividades.order_by(Atividade.data).all()
    
    # 5. Buscar Presenças (Notas)
    # Busca notas apenas dos alunos e atividades filtradas para otimizar
    presencas_db = []
    if alunos and atividades:
        ids_atividades = [a.id for a in atividades]
        ids_alunos = [a.id for a in alunos]
        
        presencas_db = Presenca.query.filter(
            Presenca.id_aluno.in_(ids_alunos),
            Presenca.id_atividade.in_(ids_atividades)
        ).all()
        
    presencas_map = { (p.id_aluno, p.id_atividade): p for p in presencas_db }

    return render_template('professor/turma/gradebook.html', 
                           turma=turma, 
                           alunos=alunos, 
                           atividades=atividades, 
                           presencas_map=presencas_map,
                           unidade_atual=unidade_selecionada, # Envia a seleção atual
                           lista_unidades=lista_unidades)     # Envia a lista completa

@alunos_bp.route('/gradebook/salvar', methods=['POST'])
@login_required
@csrf.exempt
def salvar_gradebook():
    """
    Salva notas via AJAX com tratamento robusto de erros e JSON responses.
    """
    try:
        data = request.json
        if not data:
            return jsonify({"status": "error", "message": "Dados não enviados"}), 400

        # Tratamento dos IDs
        try:
            id_aluno = int(data.get('id_aluno'))
            id_atividade = int(data.get('id_atividade'))
        except (ValueError, TypeError):
            return jsonify({"status": "error", "message": "IDs inválidos"}), 400

        campo = data.get('campo') 
        valor = data.get('valor')

        # 1. Carregar Objetos do Banco (Usando .get para evitar 404 HTML)
        aluno = Aluno.query.get(id_aluno)
        if not aluno:
            return jsonify({"status": "error", "message": "Aluno não encontrado"}), 404
            
        # Proteção Aluno Órfão
        if not aluno.turma:
            return jsonify({"status": "error", "message": "Este aluno não pertence a nenhuma turma"}), 400
            
        # Proteção Autorização
        if not aluno.turma.autor or aluno.turma.autor.id != current_user.id:
            return jsonify({"status": "error", "message": "Não autorizado"}), 403
            
        atividade = Atividade.query.get(id_atividade)
        if not atividade:
            return jsonify({"status": "error", "message": "Atividade não encontrada"}), 404

        # 2. Buscar ou Criar Presença
        presenca = Presenca.query.filter_by(id_aluno=id_aluno, id_atividade=id_atividade).first()
        
        if not presenca:
            presenca = Presenca(id_aluno=id_aluno, id_atividade=id_atividade,
                                status='Presente', participacao='Sim', situacao='Bom',
                                nota=0.0, desempenho=0)
            db.session.add(presenca)

        # 3. Atualizar Dados
        if campo == 'nota':
            if valor is not None and str(valor).strip() != '':
                try:
                    valor_str = str(valor).replace(',', '.')
                    valor_float = float(valor_str)
                    
                    if atividade.peso is not None and valor_float > atividade.peso: 
                        return jsonify({"status": "error", "message": f"Nota excede o máximo ({atividade.peso})"}), 400
                    
                    presenca.nota = valor_float
                except ValueError:
                    return jsonify({"status": "error", "message": "Valor de nota inválido"}), 400
            else:
                presenca.nota = 0.0
                
        elif campo == 'desempenho':
            try:
                presenca.desempenho = int(valor) if valor else 0
            except ValueError:
                return jsonify({"status": "error", "message": "Desempenho deve ser número inteiro"}), 400
                
        elif campo == 'status':
            presenca.status = valor
        elif campo == 'situacao':
            presenca.situacao = valor
        else:
            return jsonify({"status": "error", "message": f"Campo '{campo}' desconhecido"}), 400
            
        db.session.commit()
        return jsonify({"status": "success", "message": "Salvo"}), 200

    except Exception as e:
        db.session.rollback()
        # Log do erro no terminal para debug
        print(f"ERRO AO SALVAR NOTA: {str(e)}")
        # Retorna erro genérico, pois o frontend já foi ajustado para ignorar o bloco .catch
        return jsonify({"status": "error", "message": f"Erro interno."}), 500

@alunos_bp.route('/gradebook/salvar_massa', methods=['POST'])
@login_required
@csrf.exempt
def salvar_gradebook_massa():
    """
    Salva uma nota para todos os alunos de uma turma em uma atividade específica (requisitado pelo usuário).
    """
    try:
        data = request.json
        id_atividade = int(data.get('id_atividade'))
        valor_nota = data.get('valor_nota')
        
        # 1. Validações iniciais e parse de nota
        try:
            # Garante que . seja usado como separador decimal
            valor_nota = float(str(valor_nota).replace(',', '.'))
        except (ValueError, TypeError):
            return jsonify({"status": "error", "message": "Nota inválida."}), 400

        # 2. Busca atividade e verifica autorização
        atividade = Atividade.query.get(id_atividade)
        if not atividade:
            return jsonify({"status": "error", "message": "Atividade não encontrada."}), 404
        if atividade.turma.autor != current_user:
            return jsonify({"status": "error", "message": "Não autorizado."}), 403

        # 3. Validação do peso
        if atividade.peso is not None and valor_nota > atividade.peso: 
            return jsonify({"status": "error", "message": f"Nota excede o máximo ({atividade.peso})"}), 400
            
        # 4. Busca todos os alunos da turma
        alunos = Aluno.query.filter_by(id_turma=atividade.id_turma).all()
        
        if not alunos:
            return jsonify({"status": "error", "message": "Turma sem alunos."}), 400

        alunos_afetados = 0
        
        # 5. Itera e atualiza/cria presenças
        for aluno in alunos:
            presenca = Presenca.query.filter_by(id_aluno=aluno.id, id_atividade=id_atividade).first()
            
            if not presenca:
                presenca = Presenca(id_aluno=aluno.id, id_atividade=id_atividade,
                                    status='Presente', participacao='Sim', situacao='Bom',
                                    desempenho=0)
                db.session.add(presenca)
                
            presenca.nota = valor_nota
            alunos_afetados += 1

        db.session.commit()
        return jsonify({"status": "success", "message": f"{alunos_afetados} notas salvas em massa."}), 200

    except Exception as e:
        db.session.rollback()
        print(f"ERRO AO SALVAR NOTAS EM MASSA: {str(e)}")
        return jsonify({"status": "error", "message": "Erro interno do servidor."}), 500

@alunos_bp.route('/turma/<int:id_turma>/editar_alunos_massa', methods=['GET'])
@login_required
def editar_alunos_massa_view(id_turma):
    """
    View para exibir o formulário de edição de alunos em massa (requisitado pelo usuário).
    """
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    alunos = Aluno.query.filter_by(id_turma=id_turma).order_by(Aluno.nome).all()
    
    return render_template('professor/turma/editar_alunos_massa.html', turma=turma, alunos=alunos)

@alunos_bp.route('/turma/<int:id_turma>/salvar_alunos_massa', methods=['POST'])
@login_required
@csrf.exempt 
def salvar_alunos_massa(id_turma):
    """
    Salva a edição em massa de alunos (requisitado pelo usuário).
    """
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        return jsonify({"status": "error", "message": "Não autorizado"}), 403

    try:
        data = request.json
        alunos_data = data.get('alunos', [])
        
        # Otimiza: Busca todos os alunos da turma
        alunos_map = {a.id: a for a in Aluno.query.filter_by(id_turma=id_turma).all()}
        
        alunos_atualizados = 0
        for aluno_data in alunos_data:
            # Garante que os IDs sejam inteiros
            try:
                aluno_id = int(aluno_data.get('id'))
            except (ValueError, TypeError):
                continue

            novo_nome = aluno_data.get('nome')
            
            if aluno_id in alunos_map and novo_nome:
                aluno = alunos_map[aluno_id]
                novo_nome_limpo = novo_nome.strip()
                
                if aluno.nome != novo_nome_limpo:
                    aluno.nome = novo_nome_limpo
                    alunos_atualizados += 1
                    
        db.session.commit()
        flash(f'{alunos_atualizados} alunos atualizados com sucesso!', 'success')
        return jsonify({"status": "success", "message": f'{alunos_atualizados} alunos atualizados.'}), 200
        
    except Exception as e:
        db.session.rollback()
        print(f"ERRO AO SALVAR ALUNOS EM MASSA: {str(e)}")
        return jsonify({"status": "error", "message": "Erro interno do servidor."}), 500


@alunos_bp.route('/dashboard/<int:id_turma>')
@login_required 
def dashboard(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    alunos = Aluno.query.filter_by(id_turma=id_turma).all()
    atividades = Atividade.query.filter_by(id_turma=id_turma).all()
    
    dados_desempenho = []
    
    if not alunos or not atividades:
        return render_template(
            'professor/dashboard/turma.html', 
            turma=turma, 
            dados_desempenho=[],
            dados_frequencia={"presente": 0, "ausente": 0, "justificado": 0},
            dados_situacao={"excelente": 0, "bom": 0, "reforco": 0, "insatisfatorio": 0},
            total_alunos=len(alunos),
            total_atividades=len(atividades),
            desempenho_medio_turma=0,
            frequencia_media=0
        )

    ids_alunos_turma = [aluno.id for aluno in alunos]

    todas_presencas_turma = Presenca.query.join(Atividade).filter(
        Presenca.id_aluno.in_(ids_alunos_turma),
        Atividade.id_turma == id_turma
    ).all()

    presencas_por_aluno = {}
    for p in todas_presencas_turma:
        if p.id_aluno not in presencas_por_aluno:
            presencas_por_aluno[p.id_aluno] = []
        presencas_por_aluno[p.id_aluno].append(p)

    for aluno in alunos:
        presencas_aluno = presencas_por_aluno.get(aluno.id, [])
        desempenho_medio = (
            sum(p.desempenho for p in presencas_aluno if p.desempenho is not None) / len(presencas_aluno)
            if presencas_aluno else 0
        )
        dados_desempenho.append({
            "aluno": aluno.nome, 
            "desempenho": desempenho_medio,
            "id_aluno": aluno.id, 
            "tem_presenca": bool(presencas_aluno)
        })
        
    desempenho_total_turma = sum(item['desempenho'] for item in dados_desempenho)
    desempenho_medio_turma = desempenho_total_turma / len(dados_desempenho) if dados_desempenho else 0

    count_presente = sum(1 for p in todas_presencas_turma if p.status == 'Presente')
    count_ausente = sum(1 for p in todas_presencas_turma if p.status == 'Ausente')
    count_justificado = sum(1 for p in todas_presencas_turma if p.status == 'Justificado')
    
    total_registros = len(todas_presencas_turma)
    frequencia_media = ((count_presente + count_justificado) / total_registros) * 100 if total_registros > 0 else 0
    
    dados_frequencia = {
        "presente": count_presente, "ausente": count_ausente, "justificado": count_justificado
    }

    count_excelente = 0; count_bom = 0; count_reforco = 0; count_insat = 0
    for item in dados_desempenho:
        d = item['desempenho']
        if d == 0 and not item['tem_presenca']: continue 
        elif d >= 80: count_excelente += 1
        elif d >= 60: count_bom += 1
        elif d >= 40: count_reforco += 1
        else: count_insat += 1

    dados_situacao = {
        "excelente": count_excelente, "bom": count_bom,
        "reforco": count_reforco, "insatisfatorio": count_insat
    }
    
    return render_template(
        'professor/dashboard/turma.html', 
        turma=turma, 
        dados_desempenho=dados_desempenho,
        dados_frequencia=dados_frequencia,
        dados_situacao=dados_situacao,
        total_alunos=len(alunos),
        total_atividades=len(atividades),
        desempenho_medio_turma=desempenho_medio_turma,
        frequencia_media=frequencia_media
    )


# ------------------- ROTA AUXILIAR DE IA PARA O FRONT-END -------------------

@alunos_bp.route('/gerar_questoes_ia', methods=['POST'])
@login_required
def gerar_questoes_ia():
    api_key = current_app.config.get('GOOGLE_API_KEY')
    if not api_key:
        return jsonify({"status": "error", "message": "API Key não configurada."}), 500

    data = request.json
    tema = data.get('tema', 'Tópico geral')
    tipo_questoes = data.get('tipo', '5 questões abertas')

    prompt = f"""
    Aja como um professor. Gere uma lista de questões para uma atividade avaliativa.
    O tema é: '{tema}'.
    Formato desejado: '{tipo_questoes}'.
    Responda APENAS com o texto das questões, formatado para ser copiado e colado (use quebras de linha \n).
    Exemplo:
    1. O que foi...?
    2. Explique...
    3. Defina...
    """

    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    payload = {"contents": [{"parts": [{"text": prompt}]}]}

    try:
        response = requests.post(url, headers=headers, data=json.dumps(payload), timeout=20)
        response.raise_for_status() 
        texto_questoes = response.json()['candidates'][0]['content']['parts'][0]['text']
        return jsonify({"status": "success", "questoes": texto_questoes.strip()})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# ------------------- ROTA DE DOWNLOAD GENÉRICO DE ARQUIVOS (uploads) -------------------

@alunos_bp.route('/uploads/<path:filename>')
@login_required
def download_material(filename):
    download_name = None
    autorizado = False
    # Define diretório padrão (pode ser imgs ou docs)
    directory = current_app.config['UPLOAD_FOLDER'] 
    
    # 1. Perfil do usuário logado (Pasta imgs)
    if current_user.foto_perfil_path == filename:
        download_name = f"perfil_{current_user.username}.{filename.split('.')[-1]}"
        directory = os.path.join(current_app.config['UPLOAD_FOLDER'], 'imgs')
        autorizado = True

    # 2. Material de plano de aula (Pasta docs)
    material = Material.query.filter_by(path_arquivo=filename).first()
    if not autorizado and material:
        if material.plano_de_aula.turma.autor == current_user:
            download_name = material.nome_arquivo
            directory = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs')
            autorizado = True
    
    # 3. Anexo de diário de bordo (Pasta docs)
    if not autorizado:
        entrada = DiarioBordo.query.filter_by(path_arquivo_anexo=filename).first()
        if entrada:
            if entrada.autor_diario == current_user:
                 download_name = entrada.nome_arquivo_anexo
                 directory = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs')
                 autorizado = True
            
    # 4. Anexo de atividade (Pasta docs)
    if not autorizado:
        atividade = Atividade.query.filter_by(path_arquivo_anexo=filename).first()
        if atividade:
            if atividade.turma.autor == current_user:
                download_name = atividade.nome_arquivo_anexo
                directory = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs')
                autorizado = True

    if not autorizado:
        flash('Ficheiro não encontrado ou acesso não autorizado.', 'danger')
        return redirect(url_for('core.index'))
        
    # Tenta servir do diretório identificado (imgs ou docs).
    # Se falhar, tenta da raiz 'uploads' (para compatibilidade com arquivos antigos)
    if os.path.exists(os.path.join(directory, filename)):
        return send_from_directory(
            directory, 
            filename, 
            as_attachment=(False if 'perfil' in str(download_name) and not request.args.get('download') else True),
            download_name=download_name
        )
    else:
        # Fallback para a raiz uploads/
        return send_from_directory(
            current_app.config['UPLOAD_FOLDER'],
            filename, 
            as_attachment=(False if 'perfil' in str(download_name) and not request.args.get('download') else True),
            download_name=download_name
        )


# ------------------- ROTA DE EXPORTAÇÃO (XLSX) ORIGINAL -------------------

@alunos_bp.route('/exportar/<int:id_turma>')
@login_required 
def exportar_relatorio(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    alunos = Aluno.query.filter_by(id_turma=id_turma).all()
    dados = []
    
    atividades_da_turma_ids = [a.id for a in turma.atividades]
    
    for aluno in alunos:
        presencas_aluno = []
        if atividades_da_turma_ids:
            presencas_aluno = Presenca.query.filter(
                Presenca.id_aluno == aluno.id, 
                Presenca.id_atividade.in_(atividades_da_turma_ids)
            ).options(joinedload(Presenca.atividade)).all()
        
        if not presencas_aluno:
             dados.append({
                "Turma": turma.nome, "Aluno": aluno.nome, "Matrícula": aluno.matricula,
                "Atividade": "N/A", "Data": "N/A", "Peso": "N/A",
                "Presença": "N/A", "Participação": "N/A", "Nota": "N/A",
                "Desempenho (%)": "N/A", "Situação": "N/A",
            })
             continue 

        for p in presencas_aluno:
            atividade = p.atividade
            dados.append({
                "Turma": turma.nome, "Aluno": aluno.nome, "Matrícula": aluno.matricula,
                "Atividade": atividade.titulo if atividade else "N/A",
                "Data": atividade.data.strftime('%d/%m/%Y') if atividade and atividade.data else "N/A",
                "Peso": atividade.peso if atividade else "N/A",
                "Presença": p.status, "Participação": p.participacao, "Nota": p.nota,
                "Desempenho (%)": p.desempenho, "Situação": p.situacao,
            })

    df = pd.DataFrame(dados)
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=f"{turma.nome}")
    output.seek(0)

    return send_file(
        output, 
        download_name=f"Relatorio_{turma.nome.replace(' ', '_')}.xlsx", 
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

# ------------------- FUNÇÃO PARA AGRUPAR DADOS POR UNIDADE -------------------

def _gerar_dados_por_unidade(turma):
    alunos = Aluno.query.filter_by(id_turma=turma.id).order_by(Aluno.nome).all()
    atividades_geral = Atividade.query.filter_by(id_turma=turma.id).order_by(Atividade.data).all()
    
    presencas_geral = Presenca.query.join(Atividade).filter(Atividade.id_turma == turma.id).all()
    presencas_map = {(p.id_aluno, p.id_atividade): p for p in presencas_geral}
    
    # Determina quais unidades têm atividades
    unidades_presentes = sorted(list(set([a.unidade for a in atividades_geral if a.unidade])))
    
    # Garante uma ordem lógica se possível
    ordem_padrao = ['1ª Unidade', '2ª Unidade', '3ª Unidade', '4ª Unidade', 'Recuperação', 'Exame Final']
    unidades_ordenadas = [u for u in ordem_padrao if u in unidades_presentes]
    # Adiciona qualquer outra unidade que não esteja na lista padrão
    for u in unidades_presentes:
        if u not in unidades_ordenadas:
            unidades_ordenadas.append(u)

    dados_unidades = {}
    tipo_map = {
        'Prova': 'PROVA', 'Atividade': 'ATIV', 'Trabalho': 'TRAB',
        'Seminario': 'SEM', 'Visto': 'VISTO', 'Participacao': 'PART'
    }

    for unidade in unidades_ordenadas:
        # Filtra atividades desta unidade
        atividades_unidade = [a for a in atividades_geral if a.unidade == unidade]
        
        cabecalhos = ["ALUNO"]
        for ativ in atividades_unidade:
            prefixo = tipo_map.get(ativ.tipo, ativ.tipo[:4].upper() if ativ.tipo else 'ATIV')
            data_str = ativ.data.strftime('%d/%m') if ativ.data else "S/D"
            cabecalhos.append(f"{prefixo} {data_str}")
        
        cabecalhos.append("TOTAL")
        cabecalhos.append("SITUAÇÃO")

        linhas = []
        for aluno in alunos:
            linha = [aluno.nome]
            total_aluno = 0.0
            
            for ativ in atividades_unidade:
                presenca = presencas_map.get((aluno.id, ativ.id))
                nota = 0.0
                if presenca and presenca.nota is not None:
                    nota = presenca.nota
                
                linha.append(nota)
                total_aluno += nota
            
            linha.append(total_aluno)
            
            # Situação por Unidade
            situacao = "APROVADO" if total_aluno >= 5.0 else "REPROVADO"
            linha.append(situacao)
            
            linhas.append(linha)
        
        dados_unidades[unidade] = {
            'atividades': atividades_unidade,
            'cabecalhos': cabecalhos,
            'linhas': linhas
        }
    
    return dados_unidades

@alunos_bp.route('/turma/<int:id_turma>/exportar_matriz_xlsx')
@login_required
def exportar_matriz_xlsx(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))

    dados_por_unidade = _gerar_dados_por_unidade(turma)
    output = BytesIO()

    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        for unidade, dados in dados_por_unidade.items():
            # Limpar nome da aba (Excel limita a 31 chars)
            sheet_name = unidade[:30]
            
            df = pd.DataFrame(dados['linhas'], columns=dados['cabecalhos'])
            df.to_excel(writer, index=False, sheet_name=sheet_name)
            
            # Formatação
            ws = writer.sheets[sheet_name]
            ws.column_dimensions['A'].width = 40
            num_cols = len(dados['cabecalhos'])
            
            # Loop para largura e cores
            for i in range(2, num_cols + 1):
                col_letter = chr(64 + i) if i <= 26 else 'AA' # Simplificado para até 26 colunas, ajustar se necessário
                if i <= 26:
                     ws.column_dimensions[col_letter].width = 15

            # Cor na Situacao (Última Coluna)
            situacao_col_idx = num_cols
            for row_idx in range(2, len(dados['linhas']) + 2):
                cell = ws.cell(row=row_idx, column=situacao_col_idx)
                if cell.value == 'APROVADO':
                    cell.font = Font(color="008000", bold=True)
                elif cell.value == 'REPROVADO':
                    cell.font = Font(color="FF0000", bold=True)

    output.seek(0)
    return send_file(
        output,
        download_name=f"Matriz_Notas_{turma.nome.replace(' ', '_')}.xlsx",
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

@alunos_bp.route('/turma/<int:id_turma>/exportar_matriz_docx')
@login_required
def exportar_matriz_docx(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))

    dados_por_unidade = _gerar_dados_por_unidade(turma)

    document = Document()
    # Configuração Paisagem
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    titulo = document.add_heading(f'TURMA {turma.nome}', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for unidade, dados in dados_por_unidade.items():
        document.add_heading(unidade, level=2)
        
        table = document.add_table(rows=1, cols=len(dados['cabecalhos']))
        table.style = 'Table Grid'

        # Cabeçalhos
        hdr_cells = table.rows[0].cells
        for i, header_text in enumerate(dados['cabecalhos']):
            hdr_cells[i].text = str(header_text)
            paragraph = hdr_cells[i].paragraphs[0]
            paragraph.runs[0].bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Linhas
        situacao_idx = len(dados['cabecalhos']) - 1

        for linha_dados in dados['linhas']:
            row_cells = table.add_row().cells
            for i, item in enumerate(linha_dados):
                if i == situacao_idx: # Coluna Situação
                    p = row_cells[i].paragraphs[0]
                    p.clear()
                    run = p.add_run(str(item))
                    run.bold = True
                    if item == 'APROVADO':
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    elif item == 'REPROVADO':
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif isinstance(item, float):
                    row_cells[i].text = f"{item:.1f}".replace('.', ',')
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    row_cells[i].text = str(item)
        
        document.add_page_break() # Quebra de página entre unidades

    f = BytesIO()
    document.save(f)
    f.seek(0)

    return send_file(
        f,
        download_name=f"Matriz_Notas_{turma.nome.replace(' ', '_')}.docx",
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@alunos_bp.route('/turma/<int:id_turma>/exportar_matriz_pdf')
@login_required
def exportar_matriz_pdf(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))

    dados_por_unidade = _gerar_dados_por_unidade(turma)

    f = BytesIO()
    doc = SimpleDocTemplate(f, pagesize=landscape(A4), rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    elements = []
    styles = getSampleStyleSheet()

    title = Paragraph(f"<b>TURMA {turma.nome} - RELATÓRIO DE NOTAS</b>", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 20))

    for unidade, dados in dados_por_unidade.items():
        # Título da Unidade
        elements.append(Paragraph(f"<b>{unidade}</b>", styles['Heading2']))
        elements.append(Spacer(1, 10))

        data_table = [dados['cabecalhos']]
        for linha in dados['linhas']:
            nova_linha = []
            for item in linha:
                if isinstance(item, float):
                    nova_linha.append(f"{item:.1f}".replace('.', ','))
                else:
                    nova_linha.append(str(item))
            data_table.append(nova_linha)

        # Larguras das colunas
        page_width = landscape(A4)[0] - 60
        col_width_aluno = 150 
        num_cols_notas = len(dados['cabecalhos']) - 1
        
        if num_cols_notas > 0:
            col_width_nota = (page_width - col_width_aluno) / num_cols_notas
            col_widths = [col_width_aluno] + [col_width_nota] * num_cols_notas
        else:
            col_widths = [col_width_aluno]

        t = Table(data_table, colWidths=col_widths)
        
        # Estilos Base
        table_styles = [
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 0), (-1, -1), 8), 
        ]

        # Estilos Condicionais (Cores)
        col_situacao_idx = len(dados['cabecalhos']) - 1
        for row_idx, row_data in enumerate(dados['linhas']):
            actual_row_idx = row_idx + 1 # +1 por causa do cabeçalho
            situacao = row_data[-1] # Último item é a situação
            
            if situacao == 'APROVADO':
                table_styles.append(('TEXTCOLOR', (col_situacao_idx, actual_row_idx), (col_situacao_idx, actual_row_idx), colors.green))
                table_styles.append(('FONTNAME', (col_situacao_idx, actual_row_idx), (col_situacao_idx, actual_row_idx), 'Helvetica-Bold'))
            elif situacao == 'REPROVADO':
                table_styles.append(('TEXTCOLOR', (col_situacao_idx, actual_row_idx), (col_situacao_idx, actual_row_idx), colors.red))
                table_styles.append(('FONTNAME', (col_situacao_idx, actual_row_idx), (col_situacao_idx, actual_row_idx), 'Helvetica-Bold'))

        t.setStyle(TableStyle(table_styles))
        elements.append(t)
        
        # Quebra de página após cada unidade
        elements.append(PageBreak())

    doc.build(elements)
    f.seek(0)

    return send_file(
        f,
        download_name=f"Matriz_Notas_{turma.nome.replace(' ', '_')}.pdf",
        as_attachment=True,
        mimetype='application/pdf'
    )

@alunos_bp.route('/aluno/<int:id_aluno>/analisar_desempenho_ia', methods=['POST'])
@login_required
def analisar_desempenho_ia(id_aluno):
    aluno = Aluno.query.get_or_404(id_aluno)
    if not aluno.turma or aluno.turma.autor != current_user:
        return jsonify({"status": "error", "message": "Não autorizado"}), 403

    api_key = current_app.config.get('GOOGLE_API_KEY')
    if not api_key:
        return jsonify({"status": "error", "message": "API Key não configurada."}), 500

    presencas = Presenca.query.filter_by(id_aluno=id_aluno)\
        .join(Atividade).order_by(Atividade.data.desc()).limit(10).all()
    
    if not presencas:
        return jsonify({"status": "error", "message": "Sem dados suficientes para análise."}), 400

    historico_texto = ""
    notas = []
    for p in presencas:
        nota = p.nota if p.nota is not None else 0.0
        notas.append(nota)
        historico_texto += f"- {p.atividade.titulo} ({p.atividade.tipo}): Nota {nota}/{p.atividade.peso} | Status: {p.status}\n"

    media_recente = sum(notas) / len(notas) if notas else 0.0

    prompt = f"""
    Aja como um coordenador pedagógico experiente. Analise o desempenho do aluno '{aluno.nome}'.
    
    DADOS DO ALUNO:
    - Média das últimas atividades: {media_recente:.1f}
    - Histórico Recente:
    {historico_texto}

    TAREFA:
    1. Identifique o principal PONTO FORTE do aluno com base nos dados.
    2. Identifique o principal PONTO DE ATENÇÃO (dificuldade recorrente, faltas ou notas baixas).
    3. Sugira 2 ações práticas e curtas para o professor ajudar este aluno.

    Responda em formato HTML simples (sem tags html/body), usando <h4> para títulos, <p> para texto e <ul>/<li> para listas. Seja direto e construtivo.
    """

    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    data = {"contents": [{"parts": [{"text": prompt}]}]}

    try:
        response = requests.post(url, headers=headers, data=json.dumps(data), timeout=30)
        response.raise_for_status()
        analise = response.json()['candidates'][0]['content']['parts'][0]['text']
        return jsonify({"status": "success", "analise": analise})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@alunos_bp.route('/listar')
@login_required
def listar_alunos():
    turmas = current_user.turmas
    alunos_list = []
    for turma in turmas:
        alunos_list.extend(turma.alunos)
        
    return render_template('admin/usuarios/listar_alunos.html', alunos=alunos_list)

@alunos_bp.route('/corrigir_resposta_ia', methods=['POST'])
@login_required
def corrigir_resposta_ia():
    api_key = current_app.config.get('GOOGLE_API_KEY')
    if not api_key:
        return jsonify({"status": "error", "message": "API Key não configurada."}), 500

    data = request.json
    questao = data.get('questao', '')
    resposta_aluno = data.get('resposta', '')
    valor_max = data.get('valor_max', 10.0)

    if not questao or not resposta_aluno:
        return jsonify({"status": "error", "message": "Questão e Resposta são obrigatórias."}), 400

    prompt = f"""
    Aja como um professor especialista corrigindo uma prova subjetiva (dissertativa).
    Sua tarefa é avaliar a resposta de um aluno para uma questão específica.

    --- DADOS DA CORREÇÃO ---
    PERGUNTA DA PROVA: "{questao}"
    VALOR MÁXIMO DA QUESTÃO: {valor_max} pontos.
    
    --- RESPOSTA DO ALUNO ---
    "{resposta_aluno}"
    -------------------------

    Avalie com rigor acadêmico, mas justo. Se a resposta estiver parcialmente correta, dê nota parcial.
    
    Responda OBRIGATORIAMENTE um objeto JSON (sem markdown, apenas o JSON puro) com este formato:
    {{
        "nota": 0.0,  // (Float com uma casa decimal, ex: 8.5)
        "feedback": "Explicação curta (max 2 frases) para o aluno sobre o motivo da nota."
    }}
    """

    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    payload = {"contents": [{"parts": [{"text": prompt}]}]}

    try:
        response = requests.post(url, headers=headers, data=json.dumps(payload), timeout=30)
        response.raise_for_status()
        
        texto_ia = response.json()['candidates'][0]['content']['parts'][0]['text']
        texto_ia = texto_ia.replace('```json', '').replace('```', '').strip()
        
        resultado = json.loads(texto_ia)
        return jsonify({"status": "success", "nota": resultado['nota'], "feedback": resultado['feedback']})
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@alunos_bp.route('/corrigir_prova_foto', methods=['POST'])
@login_required
def corrigir_prova_foto():
    api_key = current_app.config.get('GOOGLE_API_KEY')
    if not api_key:
        return jsonify({"status": "error", "message": "API Key não configurada."}), 500

    # Recebe a imagem e dados
    arquivo = request.files.get('imagem_prova')
    gabarito_ou_contexto = request.form.get('contexto', '')
    valor_total = request.form.get('valor_total', 10.0)

    if not arquivo:
        return jsonify({"status": "error", "message": "Nenhuma imagem enviada."}), 400

    # Converte imagem para Base64
    imagem_base64 = base64.b64encode(arquivo.read()).decode('utf-8')
    mime_type = arquivo.mimetype # ex: image/jpeg

    prompt = f"""
    Aja como um professor corrigindo uma prova real baseada nesta imagem.
    
    DADOS DA PROVA:
    - Valor Total da Prova: {valor_total} pontos.
    - Gabarito/Contexto (Opcional): "{gabarito_ou_contexto}"

    TAREFA:
    1. Identifique as questões na imagem (Múltipla escolha e Dissertativas).
    2. Verifique as respostas do aluno (marcadas ou escritas).
    3. Corrija cada questão. Se for subjetiva, avalie a coerência.
    4. Calcule a NOTA FINAL somando os acertos.

    Retorne APENAS um JSON (sem markdown) neste formato:
    {{
        "nota_calculada": 0.0,
        "resumo_correcao": "Q1: Correta. Q2: Errou (marcou B, era C). Q3: Parcial (esqueceu X).",
        "feedback_geral": "Bom trabalho, mas atenção em..."
    }}
    """

    # CORREÇÃO: USANDO MODELO 2.5 COMO SOLICITADO
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    
    # Payload específico para envio de imagem (Multimodal)
    data = {
        "contents": [{
            "parts": [
                {"text": prompt},
                {"inline_data": {"mime_type": mime_type, "data": imagem_base64}}
            ]
        }]
    }

    try:
        response = requests.post(url, headers=headers, data=json.dumps(data), timeout=60)
        response.raise_for_status()
        
        texto_ia = response.json()['candidates'][0]['content']['parts'][0]['text']
        texto_ia = texto_ia.replace('```json', '').replace('```', '').strip()
        resultado = json.loads(texto_ia)
        
        return jsonify({"status": "success", **resultado})
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500