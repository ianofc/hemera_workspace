# blueprints/alunos.py

import os
import requests 
import json 
import base64    
from datetime import date, datetime 
from io import BytesIO

from flask import (
    Blueprint, render_template, redirect, url_for, 
    send_file, flash, jsonify, send_from_directory, request, current_app
)
import pandas as pd
from werkzeug.utils import secure_filename
from sqlalchemy import func, distinct, case
from sqlalchemy.orm import joinedload

# --- Imports para Exportação de Documentos (ReportLab e Docx) ---
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# Imports Locais e Extensões
from extensions import csrf 
from models import (
    db, Turma, Aluno, Atividade, Presenca, DiarioBordo, Material, BlocoAula, Horario
)
from forms import (
    AlunoForm, AtividadeForm, PresencaForm, EditarAlunoForm
)
from utils import extrair_texto_de_ficheiro, obter_resumo_ia, allowed_file 
from flask_login import login_required, current_user

# Criação do Blueprint
alunos_bp = Blueprint('alunos', __name__, url_prefix='/')


# ------------------- ROTAS DE TURMAS, ALUNOS E ATIVIDADES (CRUD) -------------------

@alunos_bp.route('/turma/<int:id_turma>')
@login_required 
def turma(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    
    if turma.id_user != current_user.id:
        flash('Acesso não autorizado a esta turma.', 'danger')
        return redirect(url_for('core.index'))
    
    q_aluno = request.args.get('q', '') 
    
    alunos_query = Aluno.query.filter_by(id_turma=id_turma)
    if q_aluno:
        alunos_query = alunos_query.filter(Aluno.nome.ilike(f'%{q_aluno}%'))
    
    alunos = alunos_query.order_by(Aluno.nome).all()
    
    # Carrega atividades ordenadas por data
    atividades = Atividade.query.filter_by(id_turma=id_turma).order_by(Atividade.data.desc()).all()
    
    # --- NOVA LÓGICA: Agrupamento por Unidade ---
    atividades_por_unidade = {}
    totais_por_unidade = {}
    ordem_unidades = [
        '1ª Unidade', 
        '2ª Unidade', 
        '3ª Unidade',
        '4ª Unidade', 
        'Recuperação',
        'Exame Final' 
    ]
    
    for a in atividades:
        # Se não tiver unidade definida, joga para a 3ª (padrão atual)
        u = a.unidade if a.unidade else '3ª Unidade'
        
        if u not in atividades_por_unidade:
            atividades_por_unidade[u] = []
            totais_por_unidade[u] = 0.0
        
        atividades_por_unidade[u].append(a)
        totais_por_unidade[u] += (a.peso or 0)
    # ---------------------------------------------

    ids_alunos = [aluno.id for aluno in alunos]
    
    presencas_turma = []
    if ids_alunos:
        # Otimização com joinedload para evitar N+1 queries
        presencas_turma = Presenca.query.join(Atividade).filter(
            Atividade.id_turma == id_turma,
            Presenca.id_aluno.in_(ids_alunos)
        ).options(joinedload(Presenca.atividade)).all() 

    # Soma o Valor Máximo de Pontos (Peso) de todas as atividades
    total_max_score = sum(a.peso for a in atividades if a.peso is not None)

    # Organiza presenças por aluno em um dicionário
    presencas_por_aluno = {}
    for p in presencas_turma:
        if p.id_aluno not in presencas_por_aluno:
            presencas_por_aluno[p.id_aluno] = []
        presencas_por_aluno[p.id_aluno].append(p)

    alunos_com_media = []
    for aluno in alunos:
        presencas_aluno = presencas_por_aluno.get(aluno.id, [])
        total_pontos_obtidos = 0
        
        for p in presencas_aluno:
            if p.nota is not None:
                total_pontos_obtidos += p.nota 
        
        media_final_pontos = total_pontos_obtidos 
            
        alunos_com_media.append({
            'aluno': aluno,
            'media': media_final_pontos
        })
        
    return render_template(
        'geral/turma.html', 
        turma=turma, 
        alunos_com_media=alunos_com_media, 
        atividades=atividades,
        atividades_por_unidade=atividades_por_unidade, # Passa o agrupamento
        totais_por_unidade=totais_por_unidade,         # Passa os totais
        ordem_unidades=ordem_unidades,                 # Passa a ordem de exibição
        q=q_aluno,
        total_max_score=total_max_score
    )

@alunos_bp.route('/aluno/<int:id_aluno>')
@login_required 
def aluno(id_aluno):
    aluno = Aluno.query.get_or_404(id_aluno)
    if not aluno.turma or aluno.turma.autor != current_user:
        flash('Acesso não autorizado a este aluno.', 'danger')
        return redirect(url_for('core.index'))
    
    # 1. Busca TODAS as atividades da turma (ordenadas pela unidade e depois data)
    atividades = Atividade.query.filter_by(id_turma=aluno.id_turma).order_by(Atividade.unidade, Atividade.data.desc()).all()
    
    # 2. Busca as presenças existentes deste aluno
    presencas = Presenca.query.filter_by(id_aluno=id_aluno).all()
    
    # 3. Cria um "Mapa" (Dicionário) para acesso rápido: { id_atividade: objeto_presenca }
    presencas_map = {p.id_atividade: p for p in presencas}
    
    # 4. Cálculos de Média
    total_pontos_obtidos = 0.0
    total_max_score = sum(a.peso for a in atividades if a.peso is not None)
    
    for p in presencas:
        if p.nota is not None:
            total_pontos_obtidos += p.nota
    
    media_final = total_pontos_obtidos
    
    return render_template(
        'geral/aluno.html', 
        aluno=aluno, 
        atividades=atividades,      # Passamos a lista completa de atividades
        presencas_map=presencas_map,# Passamos o mapa de presenças
        media_final=media_final,
        total_max_score=total_max_score
    )

@alunos_bp.route('/add_aluno/<int:id_turma>', methods=['GET', 'POST'])
@login_required 
def add_aluno(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    form = AlunoForm()
    if form.validate_on_submit():
        novo = Aluno(
            nome=form.nome.data, 
            matricula=form.matricula.data, 
            id_turma=id_turma, 
            data_cadastro=date.today()
        )
        db.session.add(novo)
        db.session.commit()
        flash(f'Aluno {novo.nome} adicionado!', 'success')
        return redirect(url_for('alunos.turma', id_turma=id_turma))
    
    return render_template('add/add_aluno.html', turma=turma, form=form)

@alunos_bp.route('/turma/<int:id_turma>/bulk_add_alunos', methods=['POST'])
@login_required
def bulk_add_alunos(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        return jsonify({"status": "error", "message": "Não autorizado"}), 403

    data = request.json
    nomes = data.get('nomes', [])

    if not nomes:
        return jsonify({"status": "error", "message": "Nenhum nome enviado"}), 400

    try:
        nomes_adicionados = 0
        for nome in nomes:
            nome_limpo = nome.strip()
            if nome_limpo: # Ignora linhas em branco
                novo_aluno = Aluno(
                    nome=nome_limpo,
                    id_turma=id_turma,
                    data_cadastro=date.today()
                )
                db.session.add(novo_aluno)
                nomes_adicionados += 1
        
        db.session.commit()
        return jsonify({"status": "success", "message": f"{nomes_adicionados} alunos adicionados."}), 200
    
    except Exception as e:
        db.session.rollback()
        print(f"Erro ao adicionar alunos em massa: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500


@alunos_bp.route('/add_atividade/<int:id_turma>', methods=['GET', 'POST'])
@login_required 
def add_atividade(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    form = AtividadeForm()
    if form.validate_on_submit():
        
        # --- Lógica de Upload do Anexo (Salva na pasta DOCS) ---
        arquivo = request.files.get('arquivo_anexo')
        nome_arquivo_anexo = None
        path_arquivo_anexo = None
        
        if arquivo and arquivo.filename != '' and allowed_file(arquivo.filename):
            # Define pasta de documentos
            docs_folder = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs')
            if not os.path.exists(docs_folder):
                os.makedirs(docs_folder) # Cria a pasta docs se não existir

            filename_seguro = secure_filename(arquivo.filename)
            _, ext_seguro = os.path.splitext(filename_seguro)
            filename_final = f"atividade_{id_turma}_{int(datetime.now().timestamp())}{ext_seguro}"
            
            # Salva em uploads/docs/
            filepath = os.path.join(docs_folder, filename_final)
            arquivo.save(filepath)
            
            nome_arquivo_anexo = arquivo.filename
            path_arquivo_anexo = filename_final
        # --- Fim da Lógica de Upload ---
        
        # Lógica para incluir número de questões na descrição
        descricao = form.descricao.data
        if form.num_questoes.data: 
            descricao = f"Nº de Questões: {form.num_questoes.data}\n\n{descricao}"

        atividade = Atividade(
            id_turma=id_turma, 
            titulo=form.titulo.data,
            unidade=form.unidade.data, # <--- SALVA A UNIDADE AQUI
            data=form.data.data, 
            peso=form.valor_total.data, 
            tipo=form.tipo.data,        
            descricao=descricao,
            nome_arquivo_anexo=nome_arquivo_anexo,
            path_arquivo_anexo=path_arquivo_anexo
        )
        
        db.session.add(atividade)
        db.session.commit()
        flash(f'Atividade criada na {atividade.unidade}!', 'success')
        return redirect(url_for('alunos.turma', id_turma=id_turma))
    
    ai_desc = request.args.get('ai_desc', None)
    if ai_desc:
        form.descricao.data = ai_desc
        
    return render_template('add/add_atividade.html', turma=turma, form=form)


@alunos_bp.route('/registrar_presenca/<int:id_aluno>/<int:id_atividade>', methods=['GET', 'POST'])
@login_required 
@csrf.exempt
def registrar_presenca(id_aluno, id_atividade):
    aluno = Aluno.query.get_or_404(id_aluno)
    atividade = Atividade.query.get_or_404(id_atividade)
    if aluno.turma.autor != current_user or atividade.turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    presenca = Presenca.query.filter_by(id_aluno=id_aluno, id_atividade=id_atividade).first()
    form = PresencaForm(obj=presenca) 
    
    if form.validate_on_submit():
        if form.nota.data is None:
             form.nota.data = 0.0

        # Lógica de Backup Backend: Se veio 'acertos' e 'total_questoes_manual' no request, recalcula
        total_questoes_manual = request.form.get('total_questoes_manual')
        if form.acertos.data is not None and total_questoes_manual and int(total_questoes_manual) > 0:
             acertos = form.acertos.data
             total = int(total_questoes_manual)
             if acertos <= total:
                 nota_calc = (acertos / total) * atividade.peso
                 form.nota.data = round(nota_calc, 2)

        if presenca:
            form.populate_obj(presenca)
            flash(f'Registro atualizado! Nota: {presenca.nota}', 'success')
        else:
            nova_presenca = Presenca(id_aluno=id_aluno, id_atividade=id_atividade)
            form.populate_obj(nova_presenca)
            db.session.add(nova_presenca)
            # CORREÇÃO: 'nova' -> 'nova_presenca'
            flash(f'Registro criado! Nota: {nova_presenca.nota}', 'success')
        
        db.session.commit()
        return redirect(url_for('alunos.turma', id_turma=aluno.id_turma))

    return render_template('geral/registrar_presenca.html', aluno=aluno, atividade=atividade, form=form)

@alunos_bp.route('/atividade/<int:id_atividade>/editar', methods=['GET', 'POST'])
@login_required
def edit_atividade(id_atividade):
    atividade = Atividade.query.get_or_404(id_atividade)
    if atividade.turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    class TempAtividade:
        def __init__(self, atividade):
            self.__dict__ = atividade.__dict__.copy()
            self.valor_total = atividade.peso
            
            num_questoes = None
            if atividade.descricao:
                try:
                    num_questoes_match = next((line.split(':')[-1].strip() for line in atividade.descricao.split('\n') if line.strip().startswith("Nº de Questões:")), None)
                    if num_questoes_match:
                         num_questoes = int(num_questoes_match)
                except:
                    num_questoes = None
            self.num_questoes = num_questoes 

    temp_atividade = TempAtividade(atividade)
    form = AtividadeForm(obj=temp_atividade)
    
    if form.validate_on_submit():
        
        # Lógica de Upload do Anexo
        arquivo = request.files.get('arquivo_anexo')
        if arquivo and arquivo.filename != '' and allowed_file(arquivo.filename):
            docs_folder = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs')
            if not os.path.exists(docs_folder):
                os.makedirs(docs_folder)

            if atividade.path_arquivo_anexo:
                # Tenta apagar arquivo antigo
                old_path = os.path.join(docs_folder, atividade.path_arquivo_anexo)
                if os.path.exists(old_path):
                    os.remove(old_path)
                else:
                    old_path_root = os.path.join(current_app.config['UPLOAD_FOLDER'], atividade.path_arquivo_anexo)
                    if os.path.exists(old_path_root):
                        os.remove(old_path_root)
            
            filename_seguro = secure_filename(arquivo.filename)
            _, ext_seguro = os.path.splitext(filename_seguro)
            filename_final = f"atividade_{atividade.id_turma}_{int(datetime.now().timestamp())}{ext_seguro}"
            
            filepath = os.path.join(docs_folder, filename_final)
            arquivo.save(filepath)
            
            atividade.nome_arquivo_anexo = arquivo.filename
            atividade.path_arquivo_anexo = filename_final
        
        form.populate_obj(atividade)
        
        atividade.peso = form.valor_total.data
        atividade.unidade = form.unidade.data # <--- ATUALIZA A UNIDADE AQUI

        descricao_form_data = form.descricao.data

        if form.num_questoes.data:
            lines = descricao_form_data.split('\n')
            
            start_index = 0
            if lines and lines[0].strip().startswith("Nº de Questões:"):
                try:
                    for i, line in enumerate(lines):
                        if i > 0 and line.strip() != '':
                            start_index = i
                            break
                        elif i == len(lines) - 1 and line.strip() == '':
                            start_index = len(lines)
                            break
                    if start_index == 0: 
                        if len(lines) >= 3 and lines[1].strip() == '':
                            start_index = 2
                        else:
                            start_index = 1
                except Exception:
                    start_index = 1

            descricao_base_limpa = "\n".join(lines[start_index:]).strip()
            atividade.descricao = f"Nº de Questões: {form.num_questoes.data}\n\n{descricao_base_limpa}"
        else:
            lines = atividade.descricao.split('\n') if atividade.descricao else []
            if lines and lines[0].strip().startswith("Nº de Questões:"):
                atividade.descricao = "\n".join(lines[2:]).strip()
            else:
                atividade.descricao = descricao_form_data.strip() 

        if not atividade.descricao.strip():
            atividade.descricao = None
        
        db.session.commit()
        flash(f'Atividade "{atividade.titulo}" atualizada com sucesso!', 'success')
        return redirect(url_for('alunos.turma', id_turma=atividade.id_turma))

    return render_template('edit/edit_atividade.html', form=form, atividade=atividade)


@alunos_bp.route('/atividade/<int:id_atividade>/deletar', methods=['POST'])
@login_required
def delete_atividade(id_atividade):
    atividade = Atividade.query.get_or_404(id_atividade)
    if atividade.turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))

    id_turma = atividade.id_turma
    try:
        if atividade.path_arquivo_anexo:
            filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs', atividade.path_arquivo_anexo)
            if os.path.exists(filepath):
                os.remove(filepath)
            else:
                filepath_root = os.path.join(current_app.config['UPLOAD_FOLDER'], atividade.path_arquivo_anexo)
                if os.path.exists(filepath_root):
                    os.remove(filepath_root)
        
        db.session.delete(atividade)
        db.session.commit()
        flash(f'Atividade "{atividade.titulo}" deletada.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erro ao deletar atividade: {e}', 'danger')

    return redirect(url_for('alunos.turma', id_turma=id_turma))


@alunos_bp.route('/aluno/<int:id_aluno>/editar', methods=['GET', 'POST'])
@login_required 
def editar_aluno(id_aluno):
    aluno = Aluno.query.get_or_404(id_aluno)
    if (not aluno.turma) or (aluno.turma.autor != current_user):
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    form = EditarAlunoForm(obj=aluno)
    from models import Turma 
    form.turma.query_factory = lambda: Turma.query.filter_by(autor=current_user).order_by(Turma.nome)
    
    if form.validate_on_submit():
        if form.turma.data and form.turma.data.autor != current_user:
            flash('Seleção de turma inválida.', 'danger')
            return redirect(url_for('alunos.editar_aluno', id_aluno=id_aluno))
            
        form.populate_obj(aluno) 
        db.session.commit()
        return redirect(url_for('alunos.aluno', id_aluno=aluno.id))
    
    return render_template('edit/edit_aluno.html', aluno=aluno, form=form)

@alunos_bp.route('/aluno/<int:id_aluno>/deletar', methods=['POST'])
@login_required 
def deletar_aluno(id_aluno):
    aluno = Aluno.query.get_or_404(id_aluno)
    if (not aluno.turma) or (aluno.turma.autor != current_user):
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    id_turma = aluno.id_turma
    db.session.delete(aluno)
    db.session.commit()
    flash(f'Aluno {aluno.nome} deletado com sucesso!', 'danger')
    return redirect(url_for('alunos.turma', id_turma=id_turma))

@alunos_bp.route('/turma/<int:id_turma>/editar', methods=['GET', 'POST'])
@login_required 
def editar_turma(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    from forms import TurmaForm 
    form = TurmaForm(obj=turma)
    if form.validate_on_submit():
        form.populate_obj(turma)
        db.session.commit()
        flash(f'Turma {turma.nome} atualizada com sucesso!', 'success')
        return redirect(url_for('alunos.turma', id_turma=turma.id))
    
    return render_template('edit/edit_turma.html', turma=turma, form=form)

@alunos_bp.route('/turma/<int:id_turma>/deletar', methods=['POST'])
@login_required 
def deletar_turma(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    Aluno.query.filter_by(id_turma=id_turma).update({'id_turma': None})
    
    user_horario_ids = [h.id for h in current_user.horarios]
    if user_horario_ids:
        BlocoAula.query.filter(
            BlocoAula.id_horario.in_(user_horario_ids),
            BlocoAula.id_turma == id_turma
        ).update({'id_turma': None, 'texto_alternativo': f'Turma Apagada ({turma.nome})'})
    
    db.session.delete(turma)
    db.session.commit()
    flash(f'Turma {turma.nome} deletada. Os alunos desta turma ficaram sem turma.', 'danger')
    return redirect(url_for('core.index'))

# ------------------- ROTAS DE GRADEBOOK E DASHBOARD POR TURMA -------------------

@alunos_bp.route('/turma/<int:id_turma>/gradebook')
@login_required
def gradebook(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Acesso não autorizado.', 'danger')
        return redirect(url_for('core.index'))
        
    alunos = Aluno.query.filter_by(id_turma=id_turma).order_by(Aluno.nome).all()
    # Ordena por Unidade e Data
    atividades = Atividade.query.filter_by(id_turma=id_turma).order_by(Atividade.unidade, Atividade.data).all()
    
    presencas_db = []
    if alunos:
        presencas_db = Presenca.query.filter(Presenca.id_aluno.in_([a.id for a in alunos])).all()
        
    presencas_map = { (p.id_aluno, p.id_atividade): p for p in presencas_db }

    return render_template('geral/gradebook.html', 
                           turma=turma, 
                           alunos=alunos, 
                           atividades=atividades, 
                           presencas_map=presencas_map)

@alunos_bp.route('/gradebook/salvar', methods=['POST'])
@login_required
@csrf.exempt
def salvar_gradebook():
    data = request.json
    try:
        id_aluno = int(data.get('id_aluno'))
        id_atividade = int(data.get('id_atividade'))
        campo = data.get('campo') 
        valor = data.get('valor')
    except:
        return jsonify({"status": "error", "message": "Dados inválidos."}), 400

    aluno = Aluno.query.get_or_404(id_aluno)
    if aluno.turma.autor != current_user:
        return jsonify({"status": "error", "message": "Não autorizado"}), 403
        
    atividade = Atividade.query.get_or_404(id_atividade)

    presenca = Presenca.query.filter_by(id_aluno=id_aluno, id_atividade=id_atividade).first()
    
    if not presenca:
        presenca = Presenca(id_aluno=id_aluno, id_atividade=id_atividade,
                            status='Presente', participacao='Sim', situacao='Bom',
                            nota=0.0, desempenho=0)
        db.session.add(presenca)

    try:
        if campo == 'nota':
            if valor is not None and str(valor).strip() != '':
                valor_str = str(valor).replace(',', '.')
                valor_float = float(valor_str)
                
                if atividade.peso is not None and valor_float > atividade.peso: 
                    return jsonify({"status": "error", "message": f"Nota excede o máximo permitido ({atividade.peso})"}), 400
                presenca.nota = valor_float
            else:
                presenca.nota = 0.0
                
        elif campo == 'desempenho':
            presenca.desempenho = int(valor) if valor else 0
        elif campo == 'status':
            presenca.status = valor
        elif campo == 'situacao':
            presenca.situacao = valor
        else:
            return jsonify({"status": "error", "message": "Campo desconhecido"}), 400
            
        db.session.commit()
        return jsonify({"status": "success", "message": "Salvo"}), 200
    except ValueError:
        return jsonify({"status": "error", "message": "Valor inválido (use números)"}), 400
    except Exception as e:
        db.session.rollback()
        return jsonify({"status": "error", "message": str(e)}), 500

@alunos_bp.route('/dashboard/<int:id_turma>')
@login_required 
def dashboard(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    alunos = Aluno.query.filter_by(id_turma=id_turma).all()
    atividades = Atividade.query.filter_by(id_turma=id_turma).all()
    
    dados_desempenho = []
    
    if not alunos or not atividades:
        return render_template(
            'geral/dashboard_turma.html', 
            turma=turma, 
            dados_desempenho=[],
            dados_frequencia={"presente": 0, "ausente": 0, "justificado": 0},
            dados_situacao={"excelente": 0, "bom": 0, "reforco": 0, "insatisfatorio": 0},
            total_alunos=len(alunos),
            total_atividades=len(atividades),
            desempenho_medio_turma=0,
            frequencia_media=0
        )

    ids_alunos_turma = [aluno.id for aluno in alunos]

    todas_presencas_turma = Presenca.query.join(Atividade).filter(
        Presenca.id_aluno.in_(ids_alunos_turma),
        Atividade.id_turma == id_turma
    ).all()

    presencas_por_aluno = {}
    for p in todas_presencas_turma:
        if p.id_aluno not in presencas_por_aluno:
            presencas_por_aluno[p.id_aluno] = []
        presencas_por_aluno[p.id_aluno].append(p)

    for aluno in alunos:
        presencas_aluno = presencas_por_aluno.get(aluno.id, [])
        desempenho_medio = (
            sum(p.desempenho for p in presencas_aluno if p.desempenho is not None) / len(presencas_aluno)
            if presencas_aluno else 0
        )
        dados_desempenho.append({
            "aluno": aluno.nome, 
            "desempenho": desempenho_medio,
            "id_aluno": aluno.id, 
            "tem_presenca": bool(presencas_aluno)
        })
        
    desempenho_total_turma = sum(item['desempenho'] for item in dados_desempenho)
    desempenho_medio_turma = desempenho_total_turma / len(dados_desempenho) if dados_desempenho else 0

    count_presente = sum(1 for p in todas_presencas_turma if p.status == 'Presente')
    count_ausente = sum(1 for p in todas_presencas_turma if p.status == 'Ausente')
    count_justificado = sum(1 for p in todas_presencas_turma if p.status == 'Justificado')
    
    total_registros = len(todas_presencas_turma)
    frequencia_media = ((count_presente + count_justificado) / total_registros) * 100 if total_registros > 0 else 0
    
    dados_frequencia = {
        "presente": count_presente, "ausente": count_ausente, "justificado": count_justificado
    }

    count_excelente = 0; count_bom = 0; count_reforco = 0; count_insat = 0
    for item in dados_desempenho:
        d = item['desempenho']
        if d == 0 and not item['tem_presenca']: continue 
        elif d >= 80: count_excelente += 1
        elif d >= 60: count_bom += 1
        elif d >= 40: count_reforco += 1
        else: count_insat += 1

    dados_situacao = {
        "excelente": count_excelente, "bom": count_bom,
        "reforco": count_reforco, "insatisfatorio": count_insat
    }
    
    return render_template(
        'geral/dashboard_turma.html', 
        turma=turma, 
        dados_desempenho=dados_desempenho,
        dados_frequencia=dados_frequencia,
        dados_situacao=dados_situacao,
        total_alunos=len(alunos),
        total_atividades=len(atividades),
        desempenho_medio_turma=desempenho_medio_turma,
        frequencia_media=frequencia_media
    )


# ------------------- ROTA AUXILIAR DE IA PARA O FRONT-END -------------------

@alunos_bp.route('/gerar_questoes_ia', methods=['POST'])
@login_required
def gerar_questoes_ia():
    api_key = current_app.config.get('GOOGLE_API_KEY')
    if not api_key:
        return jsonify({"status": "error", "message": "API Key não configurada."}), 500

    data = request.json
    tema = data.get('tema', 'Tópico geral')
    tipo_questoes = data.get('tipo', '5 questões abertas')

    prompt = f"""
    Aja como um professor. Gere uma lista de questões para uma atividade avaliativa.
    O tema é: '{tema}'.
    Formato desejado: '{tipo_questoes}'.
    Responda APENAS com o texto das questões, formatado para ser copiado e colado (use quebras de linha \n).
    Exemplo:
    1. O que foi...?
    2. Explique...
    3. Defina...
    """

    # CORREÇÃO: URL 2.5 Flash Preview
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    payload = {"contents": [{"parts": [{"text": prompt}]}]}

    try:
        response = requests.post(url, headers=headers, data=json.dumps(payload), timeout=20)
        response.raise_for_status() 
        texto_questoes = response.json()['candidates'][0]['content']['parts'][0]['text']
        return jsonify({"status": "success", "questoes": texto_questoes.strip()})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# ------------------- ROTA DE DOWNLOAD GENÉRICO DE ARQUIVOS (uploads) -------------------

@alunos_bp.route('/uploads/<path:filename>')
@login_required
def download_material(filename):
    download_name = None
    autorizado = False
    # Define diretório padrão (pode ser imgs ou docs)
    directory = current_app.config['UPLOAD_FOLDER'] 
    
    # 1. Perfil do usuário logado (Pasta imgs)
    if current_user.foto_perfil_path == filename:
        download_name = f"perfil_{current_user.username}.{filename.split('.')[-1]}"
        directory = os.path.join(current_app.config['UPLOAD_FOLDER'], 'imgs')
        autorizado = True

    # 2. Material de plano de aula (Pasta docs)
    material = Material.query.filter_by(path_arquivo=filename).first()
    if not autorizado and material:
        if material.plano_de_aula.turma.autor == current_user:
            download_name = material.nome_arquivo
            directory = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs')
            autorizado = True
    
    # 3. Anexo de diário de bordo (Pasta docs)
    if not autorizado:
        entrada = DiarioBordo.query.filter_by(path_arquivo_anexo=filename).first()
        if entrada:
            if entrada.autor_diario == current_user:
                 download_name = entrada.nome_arquivo_anexo
                 directory = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs')
                 autorizado = True
            
    # 4. Anexo de atividade (Pasta docs)
    if not autorizado:
        atividade = Atividade.query.filter_by(path_arquivo_anexo=filename).first()
        if atividade:
            if atividade.turma.autor == current_user:
                download_name = atividade.nome_arquivo_anexo
                directory = os.path.join(current_app.config['UPLOAD_FOLDER'], 'docs')
                autorizado = True

    if not autorizado:
        flash('Ficheiro não encontrado ou acesso não autorizado.', 'danger')
        return redirect(url_for('core.index'))
        
    # Tenta servir do diretório identificado (imgs ou docs).
    # Se falhar, tenta da raiz 'uploads' (para compatibilidade com arquivos antigos)
    if os.path.exists(os.path.join(directory, filename)):
        return send_from_directory(
            directory, 
            filename, 
            as_attachment=(False if 'perfil' in str(download_name) and not request.args.get('download') else True),
            download_name=download_name
        )
    else:
        # Fallback para a raiz uploads/
        return send_from_directory(
            current_app.config['UPLOAD_FOLDER'],
            filename, 
            as_attachment=(False if 'perfil' in str(download_name) and not request.args.get('download') else True),
            download_name=download_name
        )


# ------------------- ROTA DE EXPORTAÇÃO (XLSX) ORIGINAL -------------------

@alunos_bp.route('/exportar/<int:id_turma>')
@login_required 
def exportar_relatorio(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))
    
    alunos = Aluno.query.filter_by(id_turma=id_turma).all()
    dados = []
    
    atividades_da_turma_ids = [a.id for a in turma.atividades]
    
    for aluno in alunos:
        presencas_aluno = []
        if atividades_da_turma_ids:
            presencas_aluno = Presenca.query.filter(
                Presenca.id_aluno == aluno.id, 
                Presenca.id_atividade.in_(atividades_da_turma_ids)
            ).options(joinedload(Presenca.atividade)).all()
        
        if not presencas_aluno:
             dados.append({
                "Turma": turma.nome, "Aluno": aluno.nome, "Matrícula": aluno.matricula,
                "Atividade": "N/A", "Data": "N/A", "Peso": "N/A",
                "Presença": "N/A", "Participação": "N/A", "Nota": "N/A",
                "Desempenho (%)": "N/A", "Situação": "N/A",
            })
             continue 

        for p in presencas_aluno:
            atividade = p.atividade
            dados.append({
                "Turma": turma.nome, "Aluno": aluno.nome, "Matrícula": aluno.matricula,
                "Atividade": atividade.titulo if atividade else "N/A",
                "Data": atividade.data.strftime('%d/%m/%Y') if atividade and atividade.data else "N/A",
                "Peso": atividade.peso if atividade else "N/A",
                "Presença": p.status, "Participação": p.participacao, "Nota": p.nota,
                "Desempenho (%)": p.desempenho, "Situação": p.situacao,
            })

    df = pd.DataFrame(dados)
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=f"{turma.nome}")
    output.seek(0)

    return send_file(
        output, 
        download_name=f"Relatorio_{turma.nome.replace(' ', '_')}.xlsx", 
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

# ------------------- NOVAS ROTAS DE EXPORTAÇÃO TIPO MATRIZ -------------------

def _gerar_dados_matriz(turma):
    alunos = Aluno.query.filter_by(id_turma=turma.id).order_by(Aluno.nome).all()
    atividades = Atividade.query.filter_by(id_turma=turma.id).order_by(Atividade.data).all()
    
    presencas = Presenca.query.join(Atividade).filter(Atividade.id_turma == turma.id).all()
    presencas_map = {(p.id_aluno, p.id_atividade): p for p in presencas}

    cabecalhos = ["ALUNO"]
    for ativ in atividades:
        header = ativ.data.strftime('%d/%m') if ativ.data else ativ.titulo[:5]
        cabecalhos.append(header)
    cabecalhos.append("TOTAL")

    linhas = []
    for aluno in alunos:
        linha = [aluno.nome]
        total_aluno = 0.0
        
        for ativ in atividades:
            presenca = presencas_map.get((aluno.id, ativ.id))
            nota = 0.0
            if presenca and presenca.nota is not None:
                nota = presenca.nota
            
            linha.append(nota)
            total_aluno += nota
        
        linha.append(total_aluno)
        linhas.append(linha)
    
    return atividades, cabecalhos, linhas

@alunos_bp.route('/turma/<int:id_turma>/exportar_matriz_xlsx')
@login_required
def exportar_matriz_xlsx(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))

    atividades, cabecalhos, linhas = _gerar_dados_matriz(turma)

    df = pd.DataFrame(linhas, columns=cabecalhos)

    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Notas')
        worksheet = writer.sheets['Notas']
        worksheet.column_dimensions['A'].width = 40
        for i in range(len(atividades) + 1):
            col_letter = chr(66 + i)
            if col_letter <= 'Z':
                worksheet.column_dimensions[col_letter].width = 10

    output.seek(0)
    return send_file(
        output,
        download_name=f"Matriz_Notas_{turma.nome.replace(' ', '_')}.xlsx",
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

@alunos_bp.route('/turma/<int:id_turma>/exportar_matriz_docx')
@login_required
def exportar_matriz_docx(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))

    atividades, cabecalhos, linhas = _gerar_dados_matriz(turma)

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    titulo = document.add_heading(f'TURMA {turma.nome}', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(rows=1, cols=len(cabecalhos))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(cabecalhos):
        hdr_cells[i].text = str(header_text)
        paragraph = hdr_cells[i].paragraphs[0]
        paragraph.runs[0].bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for linha_dados in linhas:
        row_cells = table.add_row().cells
        for i, item in enumerate(linha_dados):
            if isinstance(item, float):
                row_cells[i].text = f"{item:.1f}".replace('.', ',')
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                row_cells[i].text = str(item)

    f = BytesIO()
    document.save(f)
    f.seek(0)

    return send_file(
        f,
        download_name=f"Matriz_Notas_{turma.nome.replace(' ', '_')}.docx",
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@alunos_bp.route('/turma/<int:id_turma>/exportar_matriz_pdf')
@login_required
def exportar_matriz_pdf(id_turma):
    turma = Turma.query.get_or_404(id_turma)
    if turma.autor != current_user:
        flash('Não autorizado.', 'danger')
        return redirect(url_for('core.index'))

    atividades, cabecalhos, linhas = _gerar_dados_matriz(turma)

    f = BytesIO()
    doc = SimpleDocTemplate(f, pagesize=landscape(A4), rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    elements = []
    styles = getSampleStyleSheet()

    title = Paragraph(f"<b>TURMA {turma.nome} - RELATÓRIO DE NOTAS</b>", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 20))

    data = [cabecalhos]
    for linha in linhas:
        nova_linha = []
        for item in linha:
            if isinstance(item, float):
                nova_linha.append(f"{item:.1f}".replace('.', ','))
            else:
                nova_linha.append(str(item))
        data.append(nova_linha)

    page_width = landscape(A4)[0] - 60
    col_width_aluno = 200
    rest_width = page_width - col_width_aluno
    num_cols_notas = len(cabecalhos) - 1
    col_width_nota = rest_width / num_cols_notas if num_cols_notas > 0 else 50

    col_widths = [col_width_aluno] + [col_width_nota] * num_cols_notas

    t = Table(data, colWidths=col_widths)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
    ]))
    elements.append(t)

    doc.build(elements)
    f.seek(0)

    return send_file(
        f,
        download_name=f"Matriz_Notas_{turma.nome.replace(' ', '_')}.pdf",
        as_attachment=True,
        mimetype='application/pdf'
    )

@alunos_bp.route('/aluno/<int:id_aluno>/analisar_desempenho_ia', methods=['POST'])
@login_required
def analisar_desempenho_ia(id_aluno):
    aluno = Aluno.query.get_or_404(id_aluno)
    if not aluno.turma or aluno.turma.autor != current_user:
        return jsonify({"status": "error", "message": "Não autorizado"}), 403

    api_key = current_app.config.get('GOOGLE_API_KEY')
    if not api_key:
        return jsonify({"status": "error", "message": "API Key não configurada."}), 500

    presencas = Presenca.query.filter_by(id_aluno=id_aluno)\
        .join(Atividade).order_by(Atividade.data.desc()).limit(10).all()
    
    if not presencas:
        return jsonify({"status": "error", "message": "Sem dados suficientes para análise."}), 400

    historico_texto = ""
    notas = []
    for p in presencas:
        nota = p.nota if p.nota is not None else 0.0
        notas.append(nota)
        historico_texto += f"- {p.atividade.titulo} ({p.atividade.tipo}): Nota {nota}/{p.atividade.peso} | Status: {p.status}\n"

    media_recente = sum(notas) / len(notas) if notas else 0.0

    prompt = f"""
    Aja como um coordenador pedagógico experiente. Analise o desempenho do aluno '{aluno.nome}'.
    
    DADOS DO ALUNO:
    - Média das últimas atividades: {media_recente:.1f}
    - Histórico Recente:
    {historico_texto}

    TAREFA:
    1. Identifique o principal PONTO FORTE do aluno com base nos dados.
    2. Identifique o principal PONTO DE ATENÇÃO (dificuldade recorrente, faltas ou notas baixas).
    3. Sugira 2 ações práticas e curtas para o professor ajudar este aluno.

    Responda em formato HTML simples (sem tags html/body), usando <h4> para títulos, <p> para texto e <ul>/<li> para listas. Seja direto e construtivo.
    """

    # CORREÇÃO: URL 2.5 Flash Preview
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    data = {"contents": [{"parts": [{"text": prompt}]}]}

    try:
        response = requests.post(url, headers=headers, data=json.dumps(data), timeout=30)
        response.raise_for_status()
        analise = response.json()['candidates'][0]['content']['parts'][0]['text']
        return jsonify({"status": "success", "analise": analise})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@alunos_bp.route('/listar')
@login_required
def listar_alunos():
    turmas = current_user.turmas
    alunos_list = []
    for turma in turmas:
        alunos_list.extend(turma.alunos)
        
    return render_template('list/listar_alunos.html', alunos=alunos_list)

@alunos_bp.route('/corrigir_resposta_ia', methods=['POST'])
@login_required
def corrigir_resposta_ia():
    api_key = current_app.config.get('GOOGLE_API_KEY')
    if not api_key:
        return jsonify({"status": "error", "message": "API Key não configurada."}), 500

    data = request.json
    questao = data.get('questao', '')
    resposta_aluno = data.get('resposta', '')
    valor_max = data.get('valor_max', 10.0)

    if not questao or not resposta_aluno:
        return jsonify({"status": "error", "message": "Questão e Resposta são obrigatórias."}), 400

    prompt = f"""
    Aja como um professor especialista corrigindo uma prova subjetiva (dissertativa).
    Sua tarefa é avaliar a resposta de um aluno para uma questão específica.

    --- DADOS DA CORREÇÃO ---
    PERGUNTA DA PROVA: "{questao}"
    VALOR MÁXIMO DA QUESTÃO: {valor_max} pontos.
    
    --- RESPOSTA DO ALUNO ---
    "{resposta_aluno}"
    -------------------------

    Avalie com rigor acadêmico, mas justo. Se a resposta estiver parcialmente correta, dê nota parcial.
    
    Responda OBRIGATORIAMENTE um objeto JSON (sem markdown, apenas o JSON puro) com este formato:
    {{
        "nota": 0.0,  // (Float com uma casa decimal, ex: 8.5)
        "feedback": "Explicação curta (max 2 frases) para o aluno sobre o motivo da nota."
    }}
    """

    # CORREÇÃO: URL 2.5 Flash Preview
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    payload = {"contents": [{"parts": [{"text": prompt}]}]}

    try:
        response = requests.post(url, headers=headers, data=json.dumps(payload), timeout=30)
        response.raise_for_status()
        
        texto_ia = response.json()['candidates'][0]['content']['parts'][0]['text']
        texto_ia = texto_ia.replace('```json', '').replace('```', '').strip()
        
        resultado = json.loads(texto_ia)
        return jsonify({"status": "success", "nota": resultado['nota'], "feedback": resultado['feedback']})
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@alunos_bp.route('/corrigir_prova_foto', methods=['POST'])
@login_required
@csrf.exempt # Mantido para segurança do upload
def corrigir_prova_foto():
    api_key = current_app.config.get('GOOGLE_API_KEY')
    if not api_key:
        return jsonify({"status": "error", "message": "API Key não configurada."}), 500

    # Recebe a imagem e dados
    arquivo = request.files.get('imagem_prova')
    gabarito_ou_contexto = request.form.get('contexto', '')
    valor_total = request.form.get('valor_total', 10.0)

    if not arquivo:
        return jsonify({"status": "error", "message": "Nenhuma imagem enviada."}), 400

    # Converte imagem para Base64
    imagem_base64 = base64.b64encode(arquivo.read()).decode('utf-8')
    mime_type = arquivo.mimetype # ex: image/jpeg

    prompt = f"""
    Aja como um professor corrigindo uma prova real baseada nesta imagem.
    
    DADOS DA PROVA:
    - Valor Total da Prova: {valor_total} pontos.
    - Gabarito/Contexto (Opcional): "{gabarito_ou_contexto}"

    TAREFA:
    1. Identifique as questões na imagem (Múltipla escolha e Dissertativas).
    2. Verifique as respostas do aluno (marcadas ou escritas).
    3. Corrija cada questão. Se for subjetiva, avalie a coerência.
    4. Calcule a NOTA FINAL somando os acertos.

    Retorne APENAS um JSON (sem markdown) neste formato:
    {{
        "nota_calculada": 0.0,
        "resumo_correcao": "Q1: Correta. Q2: Errou (marcou B, era C). Q3: Parcial (esqueceu X).",
        "feedback_geral": "Bom trabalho, mas atenção em..."
    }}
    """

    # CORREÇÃO: URL 2.5 Flash Preview (Uniformizada)
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    
    # Payload específico para envio de imagem (Multimodal)
    data = {
        "contents": [{
            "parts": [
                {"text": prompt},
                {"inline_data": {"mime_type": mime_type, "data": imagem_base64}}
            ]
        }]
    }

    try:
        response = requests.post(url, headers=headers, data=json.dumps(data), timeout=60)
        response.raise_for_status()
        
        texto_ia = response.json()['candidates'][0]['content']['parts'][0]['text']
        texto_ia = texto_ia.replace('```json', '').replace('```', '').strip()
        resultado = json.loads(texto_ia)
        
        return jsonify({"status": "success", **resultado})
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500